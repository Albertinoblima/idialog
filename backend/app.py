import base64
import datetime as dt
import io
import json
import os

import re
import urllib.error
import urllib.parse
import urllib.request
import uuid
from functools import wraps
from pathlib import Path

# ── Database driver (PostgreSQL or SQLite fallback) ────────────────────────
DATABASE_URL = os.getenv('DATABASE_URL', '')
USE_POSTGRES = bool(DATABASE_URL)

if USE_POSTGRES:
    import psycopg2
    import psycopg2.extras
else:
    import sqlite3

try:
    from google.analytics.data_v1beta import BetaAnalyticsDataClient
    from google.analytics.data_v1beta.types import (
        DateRange, Dimension, Metric, OrderBy, RunReportRequest,
    )
    from google.oauth2.service_account import Credentials as ServiceAccountCredentials
    GA4_AVAILABLE = True
except Exception:
    GA4_AVAILABLE = False

try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    LIMITER_AVAILABLE = True
except Exception:
    LIMITER_AVAILABLE = False

try:
    from cryptography.fernet import Fernet
    import hashlib
    CRYPTO_AVAILABLE = True
except Exception:
    CRYPTO_AVAILABLE = False

from docx import Document
from docx.shared import Inches
from flask import Flask, jsonify, request, send_file, send_from_directory
from flask_cors import CORS
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from openpyxl import Workbook, load_workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

ROOT_DIR = Path(__file__).resolve().parent
DB_PATH = ROOT_DIR / 'idialog_tools.db'
UPLOADS_DIR = ROOT_DIR / 'uploads'
UPLOADS_DIR.mkdir(exist_ok=True)

DEFAULT_TEMPLATE_PATH = r"C:\Users\alber\Downloads\PlanejamentoEstrategico iDialog - A3P Holding.xlsm"
TEMPLATE_PATH = os.getenv('IDIALOG_PLANNING_TEMPLATE', DEFAULT_TEMPLATE_PATH)

ALLOWED_LOGO_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp'}
ALLOWED_WIDGET_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp', 'gif'}

DEFAULT_ADMIN_COMPANY = os.getenv('IDIALOG_ADMIN_COMPANY', 'iDialog')
DEFAULT_ADMIN_NAME = os.getenv('IDIALOG_ADMIN_NAME', 'Albertino')
DEFAULT_ADMIN_EMAIL = os.getenv('IDIALOG_ADMIN_EMAIL', 'albertinoblima@gmail.com').strip().lower()
DEFAULT_ADMIN_PASSWORD = os.getenv('IDIALOG_ADMIN_PASSWORD', '315218')

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024
app.config['SECRET_KEY'] = os.getenv('IDIALOG_SECRET_KEY', 'idialog-dev-secret-key')

CORS(app, resources={r"/api/*": {"origins": "*"}, r"/uploads/*": {"origins": "*"}})

if LIMITER_AVAILABLE:
    limiter = Limiter(get_remote_address, app=app, default_limits=[], storage_uri='memory://')
else:
    limiter = None

def _get_fernet():
    if not CRYPTO_AVAILABLE:
        return None
    raw = app.config['SECRET_KEY'].encode()
    key = base64.urlsafe_b64encode(hashlib.sha256(raw).digest())
    return Fernet(key)

def _encrypt(value: str) -> str:
    f = _get_fernet()
    if not f or not value:
        return value
    return f.encrypt(value.encode()).decode()

def _decrypt(value: str) -> str:
    f = _get_fernet()
    if not f or not value:
        return value
    try:
        return f.decrypt(value.encode()).decode()
    except Exception:
        return value

serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])

# ── DB connection abstraction ─────────────────────────────────────────────

class _PgRow(dict):
    """Dict subclass that allows attribute-style and index access (mimics sqlite3.Row)."""
    def __getitem__(self, key):
        if isinstance(key, int):
            return list(self.values())[key]
        return super().__getitem__(key)

    def get(self, key, default=None):
        return super().get(key, default)


class _PgCursor:
    """Wraps psycopg2 cursor to auto-convert ? -> %s and return _PgRow objects."""
    def __init__(self, cur):
        self._cur = cur

    @staticmethod
    def _adapt(sql):
        return sql.replace('?', '%s')

    def execute(self, sql, params=()):
        self._cur.execute(self._adapt(sql), params)
        return self

    def executemany(self, sql, seq):
        self._cur.executemany(self._adapt(sql), seq)
        return self

    def fetchone(self):
        row = self._cur.fetchone()
        return _PgRow(row) if row else None

    def fetchall(self):
        return [_PgRow(r) for r in self._cur.fetchall()]

    @property
    def lastrowid(self):
        self._cur.execute('SELECT lastval()')
        return self._cur.fetchone()[0]

    @property
    def rowcount(self):
        return self._cur.rowcount


class _PgConn:
    """Wraps psycopg2 connection to expose sqlite3-compatible interface."""
    def __init__(self, conn):
        self._conn = conn
        self._cur = _PgCursor(conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor))

    def cursor(self):
        return _PgCursor(self._conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor))

    def execute(self, sql, params=()):
        return self._cur.execute(sql, params)

    def fetchone(self):
        return self._cur.fetchone()

    def fetchall(self):
        return self._cur.fetchall()

    def commit(self):
        self._conn.commit()

    def close(self):
        self._conn.close()


def get_db():
    if USE_POSTGRES:
        url = DATABASE_URL
        # Railway uses postgres:// but psycopg2 needs postgresql://
        if url.startswith('postgres://'):
            url = 'postgresql://' + url[len('postgres://'):]
        raw = psycopg2.connect(url)
        return _PgConn(raw)
    else:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        return conn


def ensure_default_site_admin(conn):
    company = conn.execute(
        'SELECT id FROM companies WHERE name = %s' if USE_POSTGRES else 'SELECT id FROM companies WHERE name = ?',
        (DEFAULT_ADMIN_COMPANY,),
    ).fetchone()

    timestamp = dt.datetime.utcnow().isoformat()

    if company:
        company_id = company['id']
    else:
        cur = conn.cursor()
        if USE_POSTGRES:
            cur.execute(
                """
                INSERT INTO companies (name, cnpj, email, phone, address, header_text, footer_text, logo_path, created_at, updated_at)
                VALUES (%s, '', %s, '', '', '', '', '', %s, %s) RETURNING id
                """,
                (DEFAULT_ADMIN_COMPANY, DEFAULT_ADMIN_EMAIL, timestamp, timestamp),
            )
            company_id = cur.fetchone()['id']
        else:
            cur.execute(
                """
                INSERT INTO companies (name, cnpj, email, phone, address, header_text, footer_text, logo_path, created_at, updated_at)
                VALUES (?, '', ?, '', '', '', '', '', ?, ?)
                """,
                (DEFAULT_ADMIN_COMPANY, DEFAULT_ADMIN_EMAIL, timestamp, timestamp),
            )
            company_id = cur.lastrowid

    user = conn.execute(
        'SELECT id FROM users WHERE email = %s' if USE_POSTGRES else 'SELECT id FROM users WHERE email = ?',
        (DEFAULT_ADMIN_EMAIL,),
    ).fetchone()

    password_hash = generate_password_hash(DEFAULT_ADMIN_PASSWORD)
    ph = '%s' if USE_POSTGRES else '?'
    if user:
        conn.execute(
            f"""
            UPDATE users
            SET company_id = {ph}, name = {ph}, password_hash = {ph}, role = 'admin'
            WHERE id = {ph}
            """,
            (company_id, DEFAULT_ADMIN_NAME, password_hash, user['id']),
        )
    else:
        conn.execute(
            f"""
            INSERT INTO users (company_id, name, email, password_hash, role, created_at)
            VALUES ({ph}, {ph}, {ph}, {ph}, 'admin', {ph})
            """,
            (company_id, DEFAULT_ADMIN_NAME, DEFAULT_ADMIN_EMAIL, password_hash, timestamp),
        )


def _pk(col='id'):
    """Returns the appropriate primary key definition for current DB engine."""
    if USE_POSTGRES:
        return f'{col} BIGSERIAL PRIMARY KEY'
    return f'{col} INTEGER PRIMARY KEY AUTOINCREMENT'


def _fk(col, ref_table, ref_col='id'):
    """Returns a FOREIGN KEY clause compatible with both engines."""
    if USE_POSTGRES:
        return f'FOREIGN KEY({col}) REFERENCES {ref_table}({ref_col})'
    return f'FOREIGN KEY({col}) REFERENCES {ref_table}({ref_col})'


def init_db():
    conn = get_db()
    cur = conn.cursor()

    pk = _pk()

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS companies (
            {pk},
            name TEXT NOT NULL,
            cnpj TEXT,
            email TEXT,
            phone TEXT,
            address TEXT,
            header_text TEXT,
            footer_text TEXT,
            logo_path TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS users (
            {pk},
            company_id BIGINT NOT NULL,
            name TEXT NOT NULL,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'admin',
            created_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS strategic_plans (
            {pk},
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            period_start TEXT,
            period_end TEXT,
            mission TEXT,
            vision TEXT,
            values_json TEXT,
            swot_strengths_json TEXT,
            swot_weaknesses_json TEXT,
            swot_opportunities_json TEXT,
            swot_threats_json TEXT,
            objectives_json TEXT,
            initiatives_json TEXT,
            indicators_json TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS swot_analyses (
            {pk},
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            analysis_date TEXT,
            business_unit TEXT,
            period_start TEXT,
            period_end TEXT,
            analysis_objective TEXT,
            strengths_json TEXT,
            weaknesses_json TEXT,
            opportunities_json TEXT,
            threats_json TEXT,
            so_strategies_json TEXT,
            wo_strategies_json TEXT,
            st_strategies_json TEXT,
            wt_strategies_json TEXT,
            priority_actions_json TEXT,
            critical_risks_json TEXT,
            executive_summary TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS business_plans (
            {pk},
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            industry TEXT,
            target_market TEXT,
            period_start TEXT,
            period_end TEXT,
            executive_summary TEXT,
            problem_statement TEXT,
            value_proposition TEXT,
            products_services_json TEXT,
            market_analysis TEXT,
            marketing_strategy TEXT,
            operational_plan TEXT,
            revenue_streams_json TEXT,
            cost_structure_json TEXT,
            investment_requirements_json TEXT,
            financial_projections_json TEXT,
            team_structure_json TEXT,
            milestones_json TEXT,
            risks_mitigation_json TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS feasibility_studies (
            {pk},
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            project_type TEXT,
            sector TEXT,
            sponsor TEXT,
            location TEXT,
            analysis_date TEXT,
            horizon_years INTEGER,
            capex_estimate REAL,
            opex_estimate REAL,
            technical_scope TEXT,
            technical_requirements_json TEXT,
            engineering_solution TEXT,
            technology_readiness TEXT,
            implementation_schedule TEXT,
            resource_plan_json TEXT,
            regulatory_licensing TEXT,
            environmental_social TEXT,
            demand_assumptions TEXT,
            pricing_model TEXT,
            revenue_assumptions_json TEXT,
            cost_assumptions_json TEXT,
            financing_structure TEXT,
            wacc_assumption TEXT,
            cash_flow_projection TEXT,
            economic_indicators_json TEXT,
            sensitivity_analysis TEXT,
            scenario_analysis TEXT,
            key_risks_json TEXT,
            risk_response_plan TEXT,
            final_recommendation TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS projects (
            {pk},
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            description TEXT,
            status TEXT NOT NULL DEFAULT 'planning',
            priority TEXT NOT NULL DEFAULT 'medium',
            manager_name TEXT,
            manager_email TEXT,
            start_date TEXT,
            end_date TEXT,
            budget REAL,
            objectives_json TEXT,
            milestones_json TEXT,
            risks_json TEXT,
            notes TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS project_tasks (
            {pk},
            project_id BIGINT NOT NULL,
            company_id BIGINT NOT NULL,
            title TEXT NOT NULL,
            description TEXT,
            milestone TEXT,
            assignee_name TEXT,
            assignee_email TEXT,
            status TEXT NOT NULL DEFAULT 'todo',
            priority TEXT NOT NULL DEFAULT 'medium',
            start_date TEXT,
            due_date TEXT,
            completed_date TEXT,
            progress INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id),
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS ad_widgets (
            {pk},
            company_id BIGINT NOT NULL,
            name TEXT NOT NULL,
            title TEXT,
            placement TEXT NOT NULL,
            scope TEXT NOT NULL DEFAULT 'all',
            widget_type TEXT NOT NULL DEFAULT 'image',
            media_path TEXT,
            target_url TEXT,
            alt_text TEXT,
            embed_code TEXT,
            display_order INTEGER NOT NULL DEFAULT 0,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(company_id) REFERENCES companies(id)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS content_items (
            {pk},
            company_id BIGINT NOT NULL,
            author_user_id BIGINT,
            content_type TEXT NOT NULL,
            title TEXT NOT NULL,
            slug TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'draft',
            context TEXT NOT NULL DEFAULT 'site',
            category TEXT,
            summary TEXT,
            body_html TEXT,
            cover_path TEXT,
            seo_title TEXT,
            seo_description TEXT,
            meta_description TEXT DEFAULT '',
            meta_keywords TEXT DEFAULT '',
            canonical_url TEXT DEFAULT '',
            extra_json TEXT,
            published_at TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            UNIQUE(company_id, slug),
            FOREIGN KEY(company_id) REFERENCES companies(id),
            FOREIGN KEY(author_user_id) REFERENCES users(id)
        )
    """)

    ensure_default_site_admin(conn)

    # ── CMS v2 tables ──
    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS analytics_config (
            {pk},
            ga4_measurement_id TEXT NOT NULL DEFAULT '',
            ga4_property_id TEXT NOT NULL DEFAULT '',
            ga4_service_account_json TEXT NOT NULL DEFAULT '',
            updated_at TEXT NOT NULL DEFAULT ''
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS page_blocks (
            {pk},
            page_id TEXT NOT NULL,
            block_key TEXT NOT NULL,
            content_html TEXT NOT NULL DEFAULT '',
            updated_at TEXT NOT NULL,
            UNIQUE(page_id, block_key)
        )
    """)

    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS github_config (
            {pk},
            github_pat TEXT NOT NULL DEFAULT '',
            repo_owner TEXT NOT NULL DEFAULT '',
            repo_name TEXT NOT NULL DEFAULT '',
            branch TEXT NOT NULL DEFAULT 'main',
            updated_at TEXT NOT NULL DEFAULT ''
        )
    """)

    conn.commit()
    conn.close()


def migrate_db():
    """Add new columns to existing tables without breaking existing data."""
    conn = get_db()
    cur = conn.cursor()

    if USE_POSTGRES:
        # PostgreSQL: use information_schema
        existing = cur.execute(
            "SELECT column_name FROM information_schema.columns WHERE table_name='content_items'"
        ).fetchall()
        existing_cols = [row['column_name'] for row in existing]
        for col_name, col_type in [
            ('meta_description', 'TEXT'),
            ('meta_keywords', 'TEXT'),
            ('canonical_url', 'TEXT'),
        ]:
            if col_name not in existing_cols:
                cur.execute(f"ALTER TABLE content_items ADD COLUMN IF NOT EXISTS {col_name} {col_type} DEFAULT ''")
    else:
        existing_cols = [row[1] for row in cur.execute('PRAGMA table_info(content_items)').fetchall()]
        for col_name, col_type in [
            ('meta_description', 'TEXT'),
            ('meta_keywords', 'TEXT'),
            ('canonical_url', 'TEXT'),
        ]:
            if col_name not in existing_cols:
                cur.execute(f"ALTER TABLE content_items ADD COLUMN {col_name} {col_type} DEFAULT ''")

    conn.commit()
    conn.close()


def migrate_scheduled_status():
    """Fix existing posts that are 'published' but have a future published_at — set them to 'scheduled'."""
    conn = get_db()
    cur = conn.cursor()
    ph = '%s' if USE_POSTGRES else '?'
    rows = cur.execute(
        "SELECT id, published_at FROM content_items WHERE status = 'published' AND published_at IS NOT NULL AND published_at <> ''"
    ).fetchall()
    fixed = 0
    for row in rows:
        try:
            pub_at = row['published_at'] if hasattr(row, 'keys') else row[1]
            item_id = row['id'] if hasattr(row, 'keys') else row[0]
            pub_dt = parse_iso_datetime(pub_at)
            if pub_dt and pub_dt > dt.datetime.now(dt.timezone.utc):
                cur.execute(
                    "UPDATE content_items SET status = " + ph + " WHERE id = " + ph,
                    ('scheduled', item_id)
                )
                fixed += 1
        except Exception:
            pass
    if fixed:
        conn.commit()
    conn.close()
    return fixed


def now_iso():
    return dt.datetime.now(dt.timezone.utc).isoformat()


def parse_iso_datetime(value):
    raw = (value or '').strip()
    if not raw:
        return None

    normalized = raw.replace('Z', '+00:00')
    parsed = dt.datetime.fromisoformat(normalized)
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=dt.timezone.utc)
    return parsed.astimezone(dt.timezone.utc)


def is_content_public_now(payload, reference_dt=None):
    status = (payload.get('status') or '').strip().lower()
    if status not in {'published', 'scheduled'}:
        return False

    published_at = (payload.get('published_at') or '').strip()
    if not published_at:
        # published without date = always visible; scheduled without date = not visible
        return status == 'published'

    try:
        publish_time = parse_iso_datetime(published_at)
    except ValueError:
        return False

    now_ref = reference_dt or dt.datetime.now(dt.timezone.utc)
    return publish_time <= now_ref


def token_for_user(user_id, company_id):
    return serializer.dumps({'user_id': user_id, 'company_id': company_id})


def decode_token(token):
    return serializer.loads(token, max_age=60 * 60 * 24 * 7)


def parse_json_list(value):
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return []


def row_to_company(row):
    data = dict(row)
    if data.get('logo_path'):
        data['logo_url'] = f"{request.host_url.rstrip('/')}/uploads/{data['logo_path']}"
    else:
        data['logo_url'] = ''
    return data


def row_to_plan(row):
    payload = dict(row)
    payload['values'] = json.loads(payload.pop('values_json') or '[]')
    payload['swot_strengths'] = json.loads(payload.pop('swot_strengths_json') or '[]')
    payload['swot_weaknesses'] = json.loads(payload.pop('swot_weaknesses_json') or '[]')
    payload['swot_opportunities'] = json.loads(payload.pop('swot_opportunities_json') or '[]')
    payload['swot_threats'] = json.loads(payload.pop('swot_threats_json') or '[]')
    payload['objectives'] = json.loads(payload.pop('objectives_json') or '[]')
    payload['initiatives'] = json.loads(payload.pop('initiatives_json') or '[]')
    payload['indicators'] = json.loads(payload.pop('indicators_json') or '[]')
    return payload


def row_to_swot(row):
    payload = dict(row)
    payload['strengths'] = json.loads(payload.pop('strengths_json') or '[]')
    payload['weaknesses'] = json.loads(payload.pop('weaknesses_json') or '[]')
    payload['opportunities'] = json.loads(payload.pop('opportunities_json') or '[]')
    payload['threats'] = json.loads(payload.pop('threats_json') or '[]')
    payload['so_strategies'] = json.loads(payload.pop('so_strategies_json') or '[]')
    payload['wo_strategies'] = json.loads(payload.pop('wo_strategies_json') or '[]')
    payload['st_strategies'] = json.loads(payload.pop('st_strategies_json') or '[]')
    payload['wt_strategies'] = json.loads(payload.pop('wt_strategies_json') or '[]')
    payload['priority_actions'] = json.loads(payload.pop('priority_actions_json') or '[]')
    payload['critical_risks'] = json.loads(payload.pop('critical_risks_json') or '[]')
    return payload


def row_to_business_plan(row):
    payload = dict(row)
    payload['products_services'] = json.loads(payload.pop('products_services_json') or '[]')
    payload['revenue_streams'] = json.loads(payload.pop('revenue_streams_json') or '[]')
    payload['cost_structure'] = json.loads(payload.pop('cost_structure_json') or '[]')
    payload['investment_requirements'] = json.loads(payload.pop('investment_requirements_json') or '[]')
    payload['financial_projections'] = json.loads(payload.pop('financial_projections_json') or '[]')
    payload['team_structure'] = json.loads(payload.pop('team_structure_json') or '[]')
    payload['milestones'] = json.loads(payload.pop('milestones_json') or '[]')
    payload['risks_mitigation'] = json.loads(payload.pop('risks_mitigation_json') or '[]')
    return payload


def row_to_feasibility_study(row):
    payload = dict(row)
    payload['technical_requirements'] = json.loads(payload.pop('technical_requirements_json') or '[]')
    payload['resource_plan'] = json.loads(payload.pop('resource_plan_json') or '[]')
    payload['revenue_assumptions'] = json.loads(payload.pop('revenue_assumptions_json') or '[]')
    payload['cost_assumptions'] = json.loads(payload.pop('cost_assumptions_json') or '[]')
    payload['economic_indicators'] = json.loads(payload.pop('economic_indicators_json') or '[]')
    payload['key_risks'] = json.loads(payload.pop('key_risks_json') or '[]')
    return payload


def row_to_project(row):
    payload = dict(row)
    payload['objectives'] = json.loads(payload.pop('objectives_json') or '[]')
    payload['milestones'] = json.loads(payload.pop('milestones_json') or '[]')
    payload['risks'] = json.loads(payload.pop('risks_json') or '[]')
    return payload


def row_to_task(row):
    return dict(row)


def row_to_widget(row):
    payload = dict(row)
    payload['is_active'] = bool(payload.get('is_active'))
    payload['media_url'] = ''
    if payload.get('media_path'):
        payload['media_url'] = f"{request.host_url.rstrip('/')}/uploads/{payload['media_path']}"
    return payload


def public_url_for_content(payload):
    slug = payload.get('slug') or ''
    content_type = payload.get('content_type') or 'page'
    if content_type == 'blog_post':
        return f"/pages/blog/post.html?slug={slug}"
    if content_type == 'simulado':
        return f"/pages/revista_concurso/simulado-admin.html?slug={slug}"
    return f"/pages/conteudo.html?slug={slug}"


def row_to_content(row):
    payload = dict(row)
    payload['extra'] = json.loads(payload.pop('extra_json') or '{}')
    payload['cover_url'] = ''
    if payload.get('cover_path'):
        payload['cover_url'] = f"{request.host_url.rstrip('/')}/uploads/{payload['cover_path']}"
    payload['preview_url'] = public_url_for_content(payload)
    return payload


def row_to_public_content(row):
    payload = row_to_content(row)
    return {
        'id': payload['id'],
        'content_type': payload['content_type'],
        'title': payload['title'],
        'slug': payload['slug'],
        'status': payload['status'],
        'context': payload['context'],
        'category': payload.get('category') or '',
        'summary': payload.get('summary') or '',
        'body_html': payload.get('body_html') or '',
        'cover_url': payload.get('cover_url') or '',
        'seo_title': payload.get('seo_title') or '',
        'seo_description': payload.get('seo_description') or '',
        'extra': payload.get('extra') or {},
        'published_at': payload.get('published_at') or '',
        'preview_url': payload.get('preview_url') or '',
    }


def row_to_public_widget(row):
    payload = row_to_widget(row)
    return {
        'id': payload['id'],
        'name': payload['name'],
        'title': payload.get('title') or payload['name'],
        'placement': payload['placement'],
        'scope': payload['scope'],
        'widget_type': payload['widget_type'],
        'media_url': payload['media_url'],
        'target_url': payload.get('target_url') or '',
        'alt_text': payload.get('alt_text') or payload.get('name') or 'Publicidade iDialog',
        'embed_code': payload.get('embed_code') or '',
        'display_order': payload.get('display_order') or 0,
    }


def auth_required(handler):
    @wraps(handler)
    def wrapper(*args, **kwargs):
        auth_header = request.headers.get('Authorization', '')
        token = auth_header.replace('Bearer', '').strip() or (request.args.get('token') or '').strip()
        if not token:
            return jsonify({'error': 'Token nao informado.'}), 401

        try:
            data = decode_token(token)
        except SignatureExpired:
            return jsonify({'error': 'Token expirado.'}), 401
        except BadSignature:
            return jsonify({'error': 'Token invalido.'}), 401

        conn = get_db()
        user = conn.execute(
            'SELECT id, company_id, email, name, role FROM users WHERE id = ?',
            (data.get('user_id'),),
        ).fetchone()
        conn.close()

        if not user:
            return jsonify({'error': 'Usuario nao encontrado.'}), 401

        request.current_user = dict(user)
        return handler(*args, **kwargs)

    return wrapper


def allowed_logo(filename):
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in ALLOWED_LOGO_EXTENSIONS


def allowed_widget_asset(filename):
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext in ALLOWED_WIDGET_EXTENSIONS


def get_plan_for_company(conn, company_id, plan_id):
    row = conn.execute(
        'SELECT * FROM strategic_plans WHERE id = ? AND company_id = ?',
        (plan_id, company_id),
    ).fetchone()
    return row


def get_swot_for_company(conn, company_id, swot_id):
    return conn.execute(
        'SELECT * FROM swot_analyses WHERE id = ? AND company_id = ?',
        (swot_id, company_id),
    ).fetchone()


def get_business_plan_for_company(conn, company_id, business_plan_id):
    return conn.execute(
        'SELECT * FROM business_plans WHERE id = ? AND company_id = ?',
        (business_plan_id, company_id),
    ).fetchone()


def get_feasibility_study_for_company(conn, company_id, feasibility_id):
    return conn.execute(
        'SELECT * FROM feasibility_studies WHERE id = ? AND company_id = ?',
        (feasibility_id, company_id),
    ).fetchone()


def get_project_for_company(conn, company_id, project_id):
    return conn.execute(
        'SELECT * FROM projects WHERE id = ? AND company_id = ?',
        (project_id, company_id),
    ).fetchone()


def get_task_for_project(conn, company_id, project_id, task_id):
    return conn.execute(
        'SELECT * FROM project_tasks WHERE id = ? AND project_id = ? AND company_id = ?',
        (task_id, project_id, company_id),
    ).fetchone()


def get_widget_for_company(conn, company_id, widget_id):
    return conn.execute(
        'SELECT * FROM ad_widgets WHERE id = ? AND company_id = ?',
        (widget_id, company_id),
    ).fetchone()


def get_content_for_company(conn, company_id, content_id):
    return conn.execute(
        'SELECT * FROM content_items WHERE id = ? AND company_id = ?',
        (content_id, company_id),
    ).fetchone()


def get_site_company_id(conn):
    user = conn.execute(
        'SELECT company_id FROM users WHERE email = ?',
        (DEFAULT_ADMIN_EMAIL,),
    ).fetchone()
    if user:
        return user['company_id']

    company = conn.execute(
        'SELECT id FROM companies WHERE name = ?',
        (DEFAULT_ADMIN_COMPANY,),
    ).fetchone()
    return company['id'] if company else None


def normalize_widget_payload(payload):
    normalized = {
        'name': (payload.get('name') or '').strip(),
        'title': (payload.get('title') or '').strip(),
        'placement': (payload.get('placement') or '').strip(),
        'scope': (payload.get('scope') or 'all').strip().lower(),
        'widget_type': (payload.get('widget_type') or 'image').strip().lower(),
        'media_path': (payload.get('media_path') or '').strip(),
        'target_url': (payload.get('target_url') or '').strip(),
        'alt_text': (payload.get('alt_text') or '').strip(),
        'embed_code': payload.get('embed_code') or '',
        'display_order': int(payload.get('display_order') or 0),
        'is_active': 1 if payload.get('is_active', True) else 0,
    }

    if normalized['scope'] not in {'all', 'site', 'blog', 'revista'}:
        normalized['scope'] = 'all'

    if normalized['placement'] not in {'top_banner', 'prefooter_banner', 'content_square', 'sidebar_square'}:
        raise ValueError('Posicionamento de widget invalido.')

    if normalized['widget_type'] not in {'image', 'code'}:
        raise ValueError('Tipo de widget invalido.')

    if not normalized['name']:
        raise ValueError('Nome do widget e obrigatorio.')

    if normalized['widget_type'] == 'code' and not normalized['embed_code'].strip():
        raise ValueError('Cole o codigo do widget para anuncios do tipo codigo.')

    if normalized['widget_type'] == 'image' and not normalized['media_path']:
        raise ValueError('Envie uma imagem ou GIF para anuncios do tipo imagem.')

    return normalized


def parse_extra_payload(value):
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError as exc:
            raise ValueError('JSON extra invalido.') from exc
    return {}


def normalize_content_payload(payload):
    content_type = (payload.get('content_type') or 'page').strip().lower()
    status = (payload.get('status') or 'draft').strip().lower()
    context = (payload.get('context') or 'site').strip().lower()
    title = (payload.get('title') or '').strip()
    slug = (payload.get('slug') or '').strip().lower()

    if content_type not in {'page', 'blog_post', 'simulado'}:
        raise ValueError('Tipo de conteudo invalido.')
    if status not in {'draft', 'published', 'scheduled', 'archived'}:
        raise ValueError('Status de conteudo invalido.')
    if context not in {'site', 'blog', 'revista'}:
        raise ValueError('Contexto de conteudo invalido.')
    if not title:
        raise ValueError('Titulo do conteudo e obrigatorio.')
    if not slug:
        raise ValueError('Slug do conteudo e obrigatorio.')
    if not all(ch.isalnum() or ch in {'-', '_'} for ch in slug):
        raise ValueError('Slug invalido. Use apenas letras, numeros, hifen e underscore.')

    extra = parse_extra_payload(payload.get('extra'))
    body_html = payload.get('body_html') or ''
    if content_type in {'page', 'blog_post'} and not body_html.strip():
        raise ValueError('O corpo HTML do conteudo e obrigatorio para paginas e postagens.')

    if content_type == 'simulado':
        questions = extra.get('questions')
        if questions is not None and not isinstance(questions, list):
            raise ValueError('O campo questions do simulado deve ser uma lista.')

    published_at_raw = (payload.get('published_at') or '').strip()
    if published_at_raw:
        try:
            published_at = parse_iso_datetime(published_at_raw).isoformat()
        except ValueError as exc:
            raise ValueError('Data/hora de publicacao invalida. Use formato ISO 8601.') from exc
    else:
        published_at = ''

    if status == 'published' and not published_at:
        published_at = dt.datetime.now(dt.timezone.utc).isoformat()

    # Auto-promote: if user set status=published but date is in the future, store as scheduled
    if status == 'published' and published_at:
        try:
            if parse_iso_datetime(published_at) > dt.datetime.now(dt.timezone.utc):
                status = 'scheduled'
        except ValueError:
            pass

    # Auto-promote: if status=scheduled but date has already passed, flip to published
    if status == 'scheduled' and published_at:
        try:
            if parse_iso_datetime(published_at) <= dt.datetime.now(dt.timezone.utc):
                status = 'published'
        except ValueError:
            pass

    return {
        'content_type': content_type,
        'title': title,
        'slug': slug,
        'status': status,
        'context': context,
        'category': (payload.get('category') or '').strip(),
        'summary': (payload.get('summary') or '').strip(),
        'body_html': body_html,
        'cover_path': (payload.get('cover_path') or '').strip(),
        'seo_title': (payload.get('seo_title') or '').strip(),
        'seo_description': (payload.get('seo_description') or '').strip(),
        'meta_description': (payload.get('meta_description') or '').strip(),
        'meta_keywords': (payload.get('meta_keywords') or '').strip(),
        'canonical_url': (payload.get('canonical_url') or '').strip(),
        'extra_json': json.dumps(extra, ensure_ascii=False),
        'published_at': published_at,
    }


@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'idialog-tools-api'})


@app.route('/api/admin/fix-scheduled', methods=['GET', 'POST'])
def fix_scheduled():
    """Manually trigger migration: set published posts with future dates to 'scheduled'.
    Accepts JWT auth OR ?secret=ADMIN_SECRET query param for one-time use."""
    secret_param = request.args.get('secret', '')
    admin_secret = os.getenv('ADMIN_SECRET', '')
    token_header = request.headers.get('Authorization', '')

    if not secret_param and not token_header:
        return jsonify({'error': 'Unauthorized'}), 401
    if secret_param:
        if not admin_secret or secret_param != admin_secret:
            return jsonify({'error': 'Invalid secret'}), 403
    else:
        # fall back to JWT auth
        from functools import wraps
        token = token_header.replace('Bearer ', '').strip()
        if not token:
            return jsonify({'error': 'Unauthorized'}), 401
        try:
            import jwt as pyjwt
            JWT_SECRET = os.getenv('JWT_SECRET', 'dev-secret')
            pyjwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        except Exception:
            return jsonify({'error': 'Invalid token'}), 401

    fixed = migrate_scheduled_status()
    return jsonify({'ok': True, 'fixed': fixed})


@app.route('/uploads/<path:filename>', methods=['GET'])
def uploaded_file(filename):
    return send_from_directory(UPLOADS_DIR, filename)


@app.route('/api/auth/register', methods=['POST'])
def register():
    payload = request.get_json(force=True, silent=True) or {}
    company = payload.get('company') or {}
    admin = payload.get('admin') or {}

    company_name = (company.get('name') or '').strip()
    admin_name = (admin.get('name') or '').strip()
    admin_email = (admin.get('email') or '').strip().lower()
    admin_password = admin.get('password') or ''

    if not company_name:
        return jsonify({'error': 'Nome da empresa e obrigatorio.'}), 400
    if not admin_name or not admin_email or len(admin_password) < 6:
        return jsonify({'error': 'Dados do administrador invalidos.'}), 400

    conn = get_db()
    existing = conn.execute('SELECT id FROM users WHERE email = ?', (admin_email,)).fetchone()
    if existing:
        conn.close()
        return jsonify({'error': 'Email de administrador ja cadastrado.'}), 409

    timestamp = now_iso()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO companies (name, cnpj, email, phone, address, header_text, footer_text, logo_path, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            company_name,
            company.get('cnpj', ''),
            company.get('email', ''),
            company.get('phone', ''),
            company.get('address', ''),
            company.get('header_text', ''),
            company.get('footer_text', ''),
            '',
            timestamp,
            timestamp,
        ),
    )
    company_id = cur.lastrowid

    cur.execute(
        """
        INSERT INTO users (company_id, name, email, password_hash, role, created_at)
        VALUES (?, ?, ?, ?, 'admin', ?)
        """,
        (
            company_id,
            admin_name,
            admin_email,
            generate_password_hash(admin_password),
            timestamp,
        ),
    )
    user_id = cur.lastrowid

    conn.commit()
    conn.close()

    return jsonify({'token': token_for_user(user_id, company_id)})


@app.route('/api/auth/login', methods=['POST'])
def login():
    payload = request.get_json(force=True, silent=True) or {}
    email = (payload.get('email') or '').strip().lower()
    password = payload.get('password') or ''

    if not email or not password:
        return jsonify({'error': 'Credenciais invalidas.'}), 400

    conn = get_db()
    user = conn.execute(
        'SELECT id, company_id, password_hash FROM users WHERE email = ?',
        (email,),
    ).fetchone()
    conn.close()

    if not user or not check_password_hash(user['password_hash'], password):
        return jsonify({'error': 'Email ou senha invalidos.'}), 401

    return jsonify({'token': token_for_user(user['id'], user['company_id'])})


@app.route('/api/me', methods=['GET'])
@auth_required
def me():
    conn = get_db()
    company_row = conn.execute(
        'SELECT * FROM companies WHERE id = ?',
        (request.current_user['company_id'],),
    ).fetchone()
    conn.close()

    if not company_row:
        return jsonify({'error': 'Empresa nao encontrada.'}), 404

    return jsonify({'user': request.current_user, 'company': row_to_company(company_row)})


@app.route('/api/company', methods=['PUT'])
@auth_required
def update_company():
    payload = request.get_json(force=True, silent=True) or {}
    name = (payload.get('name') or '').strip()

    if not name:
        return jsonify({'error': 'Nome da empresa e obrigatorio.'}), 400

    timestamp = now_iso()
    conn = get_db()
    conn.execute(
        """
        UPDATE companies
        SET name = ?, cnpj = ?, email = ?, phone = ?, address = ?, header_text = ?, footer_text = ?, updated_at = ?
        WHERE id = ?
        """,
        (
            name,
            payload.get('cnpj', ''),
            payload.get('email', ''),
            payload.get('phone', ''),
            payload.get('address', ''),
            payload.get('header_text', ''),
            payload.get('footer_text', ''),
            timestamp,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    company_row = conn.execute(
        'SELECT * FROM companies WHERE id = ?',
        (request.current_user['company_id'],),
    ).fetchone()
    conn.close()

    return jsonify({'company': row_to_company(company_row)})


@app.route('/api/company/logo', methods=['POST'])
@auth_required
def upload_logo():
    if 'logo' not in request.files:
        return jsonify({'error': 'Arquivo de logo nao enviado.'}), 400

    logo = request.files['logo']
    if not logo or not logo.filename:
        return jsonify({'error': 'Logo invalida.'}), 400

    if not allowed_logo(logo.filename):
        return jsonify({'error': 'Formato de logo nao permitido.'}), 400

    ext = logo.filename.rsplit('.', 1)[1].lower()
    filename = secure_filename(f"company_{request.current_user['company_id']}_{uuid.uuid4().hex}.{ext}")
    filepath = UPLOADS_DIR / filename
    logo.save(filepath)

    conn = get_db()
    conn.execute(
        'UPDATE companies SET logo_path = ?, updated_at = ? WHERE id = ?',
        (filename, now_iso(), request.current_user['company_id']),
    )
    conn.commit()
    company_row = conn.execute(
        'SELECT * FROM companies WHERE id = ?',
        (request.current_user['company_id'],),
    ).fetchone()
    conn.close()

    return jsonify({'company': row_to_company(company_row)})


@app.route('/api/widgets/media', methods=['POST'])
@auth_required
def upload_widget_media():
    if 'media' not in request.files:
        return jsonify({'error': 'Arquivo do widget nao enviado.'}), 400

    media = request.files['media']
    if not media or not media.filename:
        return jsonify({'error': 'Arquivo do widget invalido.'}), 400

    if not allowed_widget_asset(media.filename):
        return jsonify({'error': 'Formato de arquivo nao permitido para widget.'}), 400

    ext = media.filename.rsplit('.', 1)[1].lower()
    filename = secure_filename(f"widget_{request.current_user['company_id']}_{uuid.uuid4().hex}.{ext}")
    filepath = UPLOADS_DIR / filename
    media.save(filepath)

    return jsonify({
        'media_path': filename,
        'media_url': f"{request.host_url.rstrip('/')}/uploads/{filename}",
    })


@app.route('/api/content/media', methods=['POST'])
@auth_required
def upload_content_media():
    if 'media' not in request.files:
        return jsonify({'error': 'Arquivo de capa nao enviado.'}), 400

    media = request.files['media']
    if not media or not media.filename:
        return jsonify({'error': 'Arquivo de capa invalido.'}), 400

    if not allowed_widget_asset(media.filename):
        return jsonify({'error': 'Formato de capa nao permitido.'}), 400

    ext = media.filename.rsplit('.', 1)[1].lower()
    filename = secure_filename(f"content_{request.current_user['company_id']}_{uuid.uuid4().hex}.{ext}")
    filepath = UPLOADS_DIR / filename
    media.save(filepath)

    return jsonify({
        'cover_path': filename,
        'cover_url': f"{request.host_url.rstrip('/')}/uploads/{filename}",
    })


@app.route('/api/content', methods=['GET'])
@auth_required
def list_content():
    content_type = (request.args.get('type') or '').strip().lower()

    query = 'SELECT * FROM content_items WHERE company_id = ?'
    params = [request.current_user['company_id']]
    if content_type:
        query += ' AND content_type = ?'
        params.append(content_type)
    query += ' ORDER BY updated_at DESC, created_at DESC'

    conn = get_db()
    rows = conn.execute(query, tuple(params)).fetchall()
    conn.close()

    return jsonify({'items': [row_to_content(row) for row in rows]})


@app.route('/api/content', methods=['POST'])
@auth_required
def create_content():
    payload = request.get_json(force=True, silent=True) or {}
    try:
        content = normalize_content_payload(payload)
    except ValueError as error:
        return jsonify({'error': str(error)}), 400

    timestamp = now_iso()
    conn = get_db()
    duplicate = conn.execute(
        'SELECT id FROM content_items WHERE company_id = ? AND slug = ?',
        (request.current_user['company_id'], content['slug']),
    ).fetchone()
    if duplicate:
        conn.close()
        return jsonify({'error': 'Ja existe um conteudo com esse slug.'}), 409

    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO content_items (
            company_id, author_user_id, content_type, title, slug, status, context,
            category, summary, body_html, cover_path, seo_title, seo_description,
            meta_description, meta_keywords, canonical_url,
            extra_json, published_at, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            request.current_user['id'],
            content['content_type'],
            content['title'],
            content['slug'],
            content['status'],
            content['context'],
            content['category'],
            content['summary'],
            content['body_html'],
            content['cover_path'],
            content['seo_title'],
            content['seo_description'],
            content['meta_description'],
            content['meta_keywords'],
            content['canonical_url'],
            content['extra_json'],
            content['published_at'],
            timestamp,
            timestamp,
        ),
    )
    item_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM content_items WHERE id = ?', (item_id,)).fetchone()
    conn.close()
    return jsonify({'item': row_to_content(row)}), 201


@app.route('/api/content/<int:content_id>', methods=['PUT'])
@auth_required
def update_content(content_id):
    payload = request.get_json(force=True, silent=True) or {}
    try:
        content = normalize_content_payload(payload)
    except ValueError as error:
        return jsonify({'error': str(error)}), 400

    conn = get_db()
    existing = get_content_for_company(conn, request.current_user['company_id'], content_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404

    duplicate = conn.execute(
        'SELECT id FROM content_items WHERE company_id = ? AND slug = ? AND id <> ?',
        (request.current_user['company_id'], content['slug'], content_id),
    ).fetchone()
    if duplicate:
        conn.close()
        return jsonify({'error': 'Ja existe outro conteudo com esse slug.'}), 409

    conn.execute(
        """
        UPDATE content_items
        SET content_type = ?, title = ?, slug = ?, status = ?, context = ?, category = ?,
            summary = ?, body_html = ?, cover_path = ?, seo_title = ?, seo_description = ?,
            meta_description = ?, meta_keywords = ?, canonical_url = ?,
            extra_json = ?, published_at = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            content['content_type'],
            content['title'],
            content['slug'],
            content['status'],
            content['context'],
            content['category'],
            content['summary'],
            content['body_html'],
            content['cover_path'],
            content['seo_title'],
            content['seo_description'],
            content['meta_description'],
            content['meta_keywords'],
            content['canonical_url'],
            content['extra_json'],
            content['published_at'],
            now_iso(),
            content_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM content_items WHERE id = ?', (content_id,)).fetchone()
    conn.close()
    return jsonify({'item': row_to_content(row)})


@app.route('/api/content/<int:content_id>', methods=['DELETE'])
@auth_required
def delete_content(content_id):
    conn = get_db()
    existing = get_content_for_company(conn, request.current_user['company_id'], content_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM content_items WHERE id = ? AND company_id = ?',
        (content_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/public/content', methods=['GET'])
def public_content_list():
    content_type = (request.args.get('type') or '').strip().lower()
    context = (request.args.get('context') or '').strip().lower()
    limit = int(request.args.get('limit') or 50)
    limit = max(1, min(limit, 200))

    conn = get_db()
    company_id = get_site_company_id(conn)
    if not company_id:
        conn.close()
        return jsonify({'items': []})

    query = 'SELECT * FROM content_items WHERE company_id = ? AND status IN (?, ?)'
    params = [company_id, 'published', 'scheduled']
    if content_type:
        query += ' AND content_type = ?'
        params.append(content_type)
    if context:
        query += ' AND context = ?'
        params.append(context)
    query += ' ORDER BY COALESCE(published_at, updated_at) DESC'

    rows = conn.execute(query, tuple(params)).fetchall()
    conn.close()

    now_ref = dt.datetime.now(dt.timezone.utc)
    visible = []
    for row in rows:
        payload = row_to_public_content(row)
        if is_content_public_now(payload, now_ref):
            visible.append(payload)
        if len(visible) >= limit:
            break

    return jsonify({'items': visible})


@app.route('/api/public/content/<slug>', methods=['GET'])
def public_content_detail(slug):
    clean_slug = (slug or '').strip().lower()
    conn = get_db()
    company_id = get_site_company_id(conn)
    if not company_id:
        conn.close()
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404

    row = conn.execute(
        'SELECT * FROM content_items WHERE company_id = ? AND slug = ? AND status IN (?, ?)',
        (company_id, clean_slug, 'published', 'scheduled'),
    ).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404

    payload = row_to_public_content(row)
    if not is_content_public_now(payload):
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404

    return jsonify({'item': payload})


@app.route('/api/widgets', methods=['GET'])
@auth_required
def list_widgets():
    conn = get_db()
    rows = conn.execute(
        """
        SELECT * FROM ad_widgets
        WHERE company_id = ?
        ORDER BY placement ASC, display_order ASC, updated_at DESC
        """,
        (request.current_user['company_id'],),
    ).fetchall()
    conn.close()

    return jsonify({'widgets': [row_to_widget(row) for row in rows]})


@app.route('/api/widgets', methods=['POST'])
@auth_required
def create_widget():
    payload = request.get_json(force=True, silent=True) or {}

    try:
        widget = normalize_widget_payload(payload)
    except ValueError as error:
        return jsonify({'error': str(error)}), 400

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO ad_widgets (
            company_id, name, title, placement, scope, widget_type, media_path,
            target_url, alt_text, embed_code, display_order, is_active, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            widget['name'],
            widget['title'],
            widget['placement'],
            widget['scope'],
            widget['widget_type'],
            widget['media_path'],
            widget['target_url'],
            widget['alt_text'],
            widget['embed_code'],
            widget['display_order'],
            widget['is_active'],
            timestamp,
            timestamp,
        ),
    )
    widget_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM ad_widgets WHERE id = ?', (widget_id,)).fetchone()
    conn.close()

    return jsonify({'widget': row_to_widget(row)}), 201


@app.route('/api/widgets/<int:widget_id>', methods=['PUT'])
@auth_required
def update_widget(widget_id):
    payload = request.get_json(force=True, silent=True) or {}

    try:
        widget = normalize_widget_payload(payload)
    except ValueError as error:
        return jsonify({'error': str(error)}), 400

    conn = get_db()
    existing = get_widget_for_company(conn, request.current_user['company_id'], widget_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Widget nao encontrado.'}), 404

    conn.execute(
        """
        UPDATE ad_widgets
        SET name = ?, title = ?, placement = ?, scope = ?, widget_type = ?, media_path = ?,
            target_url = ?, alt_text = ?, embed_code = ?, display_order = ?, is_active = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            widget['name'],
            widget['title'],
            widget['placement'],
            widget['scope'],
            widget['widget_type'],
            widget['media_path'],
            widget['target_url'],
            widget['alt_text'],
            widget['embed_code'],
            widget['display_order'],
            widget['is_active'],
            now_iso(),
            widget_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM ad_widgets WHERE id = ?', (widget_id,)).fetchone()
    conn.close()

    return jsonify({'widget': row_to_widget(row)})


@app.route('/api/widgets/<int:widget_id>', methods=['DELETE'])
@auth_required
def delete_widget(widget_id):
    conn = get_db()
    existing = get_widget_for_company(conn, request.current_user['company_id'], widget_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Widget nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM ad_widgets WHERE id = ? AND company_id = ?',
        (widget_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()

    return jsonify({'ok': True})


@app.route('/api/public/widgets', methods=['GET'])
def public_widgets():
    context = (request.args.get('context') or 'site').strip().lower()
    if context not in {'site', 'blog', 'revista'}:
        context = 'site'

    scope_map = {
        'site': ('all', 'site'),
        'blog': ('all', 'blog'),
        'revista': ('all', 'revista'),
    }

    conn = get_db()
    company_id = get_site_company_id(conn)
    if not company_id:
        conn.close()
        return jsonify({'context': context, 'placements': {}})

    rows = conn.execute(
        """
        SELECT * FROM ad_widgets
        WHERE company_id = ? AND is_active = 1 AND scope IN (?, ?)
        ORDER BY placement ASC, display_order ASC, updated_at DESC
        """,
        (company_id, scope_map[context][0], scope_map[context][1]),
    ).fetchall()
    conn.close()

    placements = {
        'top_banner': [],
        'prefooter_banner': [],
        'content_square': [],
        'sidebar_square': [],
    }
    for row in rows:
        widget = row_to_public_widget(row)
        placements[widget['placement']].append(widget)

    return jsonify({'context': context, 'placements': placements})


@app.route('/api/swot-analyses', methods=['GET'])
@auth_required
def list_swot_analyses():
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM swot_analyses WHERE company_id = ? ORDER BY updated_at DESC',
        (request.current_user['company_id'],),
    ).fetchall()
    conn.close()
    return jsonify({'analyses': [row_to_swot(row) for row in rows]})


@app.route('/api/swot-analyses', methods=['POST'])
@auth_required
def create_swot_analysis():
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo da analise SWOT e obrigatorio.'}), 400

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO swot_analyses (
            company_id, title, analysis_date, business_unit, period_start, period_end,
            analysis_objective, strengths_json, weaknesses_json, opportunities_json,
            threats_json, so_strategies_json, wo_strategies_json, st_strategies_json,
            wt_strategies_json, priority_actions_json, critical_risks_json,
            executive_summary, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            title,
            payload.get('analysis_date', ''),
            payload.get('business_unit', ''),
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('analysis_objective', ''),
            json.dumps(parse_json_list(payload.get('strengths'))),
            json.dumps(parse_json_list(payload.get('weaknesses'))),
            json.dumps(parse_json_list(payload.get('opportunities'))),
            json.dumps(parse_json_list(payload.get('threats'))),
            json.dumps(parse_json_list(payload.get('so_strategies'))),
            json.dumps(parse_json_list(payload.get('wo_strategies'))),
            json.dumps(parse_json_list(payload.get('st_strategies'))),
            json.dumps(parse_json_list(payload.get('wt_strategies'))),
            json.dumps(parse_json_list(payload.get('priority_actions'))),
            json.dumps(parse_json_list(payload.get('critical_risks'))),
            payload.get('executive_summary', ''),
            timestamp,
            timestamp,
        ),
    )
    swot_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM swot_analyses WHERE id = ?', (swot_id,)).fetchone()
    conn.close()
    return jsonify({'analysis': row_to_swot(row)}), 201


@app.route('/api/swot-analyses/<int:swot_id>', methods=['PUT'])
@auth_required
def update_swot_analysis(swot_id):
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo da analise SWOT e obrigatorio.'}), 400

    conn = get_db()
    existing = get_swot_for_company(conn, request.current_user['company_id'], swot_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Analise SWOT nao encontrada.'}), 404

    conn.execute(
        """
        UPDATE swot_analyses
        SET title = ?, analysis_date = ?, business_unit = ?, period_start = ?, period_end = ?,
            analysis_objective = ?, strengths_json = ?, weaknesses_json = ?, opportunities_json = ?,
            threats_json = ?, so_strategies_json = ?, wo_strategies_json = ?, st_strategies_json = ?,
            wt_strategies_json = ?, priority_actions_json = ?, critical_risks_json = ?,
            executive_summary = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('analysis_date', ''),
            payload.get('business_unit', ''),
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('analysis_objective', ''),
            json.dumps(parse_json_list(payload.get('strengths'))),
            json.dumps(parse_json_list(payload.get('weaknesses'))),
            json.dumps(parse_json_list(payload.get('opportunities'))),
            json.dumps(parse_json_list(payload.get('threats'))),
            json.dumps(parse_json_list(payload.get('so_strategies'))),
            json.dumps(parse_json_list(payload.get('wo_strategies'))),
            json.dumps(parse_json_list(payload.get('st_strategies'))),
            json.dumps(parse_json_list(payload.get('wt_strategies'))),
            json.dumps(parse_json_list(payload.get('priority_actions'))),
            json.dumps(parse_json_list(payload.get('critical_risks'))),
            payload.get('executive_summary', ''),
            now_iso(),
            swot_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM swot_analyses WHERE id = ?', (swot_id,)).fetchone()
    conn.close()
    return jsonify({'analysis': row_to_swot(row)})


@app.route('/api/swot-analyses/<int:swot_id>', methods=['DELETE'])
@auth_required
def delete_swot_analysis(swot_id):
    conn = get_db()
    existing = get_swot_for_company(conn, request.current_user['company_id'], swot_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Analise SWOT nao encontrada.'}), 404

    conn.execute(
        'DELETE FROM swot_analyses WHERE id = ? AND company_id = ?',
        (swot_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/business-plans', methods=['GET'])
@auth_required
def list_business_plans():
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM business_plans WHERE company_id = ? ORDER BY updated_at DESC',
        (request.current_user['company_id'],),
    ).fetchall()
    conn.close()
    return jsonify({'business_plans': [row_to_business_plan(row) for row in rows]})


@app.route('/api/business-plans', methods=['POST'])
@auth_required
def create_business_plan():
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do plano de negocios e obrigatorio.'}), 400

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO business_plans (
            company_id, title, industry, target_market, period_start, period_end,
            executive_summary, problem_statement, value_proposition, products_services_json,
            market_analysis, marketing_strategy, operational_plan, revenue_streams_json,
            cost_structure_json, investment_requirements_json, financial_projections_json,
            team_structure_json, milestones_json, risks_mitigation_json, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            title,
            payload.get('industry', ''),
            payload.get('target_market', ''),
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('executive_summary', ''),
            payload.get('problem_statement', ''),
            payload.get('value_proposition', ''),
            json.dumps(parse_json_list(payload.get('products_services'))),
            payload.get('market_analysis', ''),
            payload.get('marketing_strategy', ''),
            payload.get('operational_plan', ''),
            json.dumps(parse_json_list(payload.get('revenue_streams'))),
            json.dumps(parse_json_list(payload.get('cost_structure'))),
            json.dumps(parse_json_list(payload.get('investment_requirements'))),
            json.dumps(parse_json_list(payload.get('financial_projections'))),
            json.dumps(parse_json_list(payload.get('team_structure'))),
            json.dumps(parse_json_list(payload.get('milestones'))),
            json.dumps(parse_json_list(payload.get('risks_mitigation'))),
            timestamp,
            timestamp,
        ),
    )
    business_plan_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM business_plans WHERE id = ?', (business_plan_id,)).fetchone()
    conn.close()
    return jsonify({'business_plan': row_to_business_plan(row)}), 201


@app.route('/api/business-plans/<int:business_plan_id>', methods=['PUT'])
@auth_required
def update_business_plan(business_plan_id):
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do plano de negocios e obrigatorio.'}), 400

    conn = get_db()
    existing = get_business_plan_for_company(conn, request.current_user['company_id'], business_plan_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Plano de negocios nao encontrado.'}), 404

    conn.execute(
        """
        UPDATE business_plans
        SET title = ?, industry = ?, target_market = ?, period_start = ?, period_end = ?,
            executive_summary = ?, problem_statement = ?, value_proposition = ?, products_services_json = ?,
            market_analysis = ?, marketing_strategy = ?, operational_plan = ?, revenue_streams_json = ?,
            cost_structure_json = ?, investment_requirements_json = ?, financial_projections_json = ?,
            team_structure_json = ?, milestones_json = ?, risks_mitigation_json = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('industry', ''),
            payload.get('target_market', ''),
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('executive_summary', ''),
            payload.get('problem_statement', ''),
            payload.get('value_proposition', ''),
            json.dumps(parse_json_list(payload.get('products_services'))),
            payload.get('market_analysis', ''),
            payload.get('marketing_strategy', ''),
            payload.get('operational_plan', ''),
            json.dumps(parse_json_list(payload.get('revenue_streams'))),
            json.dumps(parse_json_list(payload.get('cost_structure'))),
            json.dumps(parse_json_list(payload.get('investment_requirements'))),
            json.dumps(parse_json_list(payload.get('financial_projections'))),
            json.dumps(parse_json_list(payload.get('team_structure'))),
            json.dumps(parse_json_list(payload.get('milestones'))),
            json.dumps(parse_json_list(payload.get('risks_mitigation'))),
            now_iso(),
            business_plan_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM business_plans WHERE id = ?', (business_plan_id,)).fetchone()
    conn.close()
    return jsonify({'business_plan': row_to_business_plan(row)})


@app.route('/api/business-plans/<int:business_plan_id>', methods=['DELETE'])
@auth_required
def delete_business_plan(business_plan_id):
    conn = get_db()
    existing = get_business_plan_for_company(conn, request.current_user['company_id'], business_plan_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Plano de negocios nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM business_plans WHERE id = ? AND company_id = ?',
        (business_plan_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/feasibility-studies', methods=['GET'])
@auth_required
def list_feasibility_studies():
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM feasibility_studies WHERE company_id = ? ORDER BY updated_at DESC',
        (request.current_user['company_id'],),
    ).fetchall()
    conn.close()
    return jsonify({'feasibility_studies': [row_to_feasibility_study(row) for row in rows]})


@app.route('/api/feasibility-studies', methods=['POST'])
@auth_required
def create_feasibility_study():
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do estudo de viabilidade e obrigatorio.'}), 400

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO feasibility_studies (
            company_id, title, project_type, sector, sponsor, location, analysis_date,
            horizon_years, capex_estimate, opex_estimate, technical_scope,
            technical_requirements_json, engineering_solution, technology_readiness,
            implementation_schedule, resource_plan_json, regulatory_licensing,
            environmental_social, demand_assumptions, pricing_model,
            revenue_assumptions_json, cost_assumptions_json, financing_structure,
            wacc_assumption, cash_flow_projection, economic_indicators_json,
            sensitivity_analysis, scenario_analysis, key_risks_json,
            risk_response_plan, final_recommendation, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            title,
            payload.get('project_type', ''),
            payload.get('sector', ''),
            payload.get('sponsor', ''),
            payload.get('location', ''),
            payload.get('analysis_date', ''),
            int(payload.get('horizon_years') or 0),
            float(payload.get('capex_estimate') or 0),
            float(payload.get('opex_estimate') or 0),
            payload.get('technical_scope', ''),
            json.dumps(parse_json_list(payload.get('technical_requirements'))),
            payload.get('engineering_solution', ''),
            payload.get('technology_readiness', ''),
            payload.get('implementation_schedule', ''),
            json.dumps(parse_json_list(payload.get('resource_plan'))),
            payload.get('regulatory_licensing', ''),
            payload.get('environmental_social', ''),
            payload.get('demand_assumptions', ''),
            payload.get('pricing_model', ''),
            json.dumps(parse_json_list(payload.get('revenue_assumptions'))),
            json.dumps(parse_json_list(payload.get('cost_assumptions'))),
            payload.get('financing_structure', ''),
            payload.get('wacc_assumption', ''),
            payload.get('cash_flow_projection', ''),
            json.dumps(parse_json_list(payload.get('economic_indicators'))),
            payload.get('sensitivity_analysis', ''),
            payload.get('scenario_analysis', ''),
            json.dumps(parse_json_list(payload.get('key_risks'))),
            payload.get('risk_response_plan', ''),
            payload.get('final_recommendation', ''),
            timestamp,
            timestamp,
        ),
    )
    feasibility_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM feasibility_studies WHERE id = ?', (feasibility_id,)).fetchone()
    conn.close()
    return jsonify({'feasibility_study': row_to_feasibility_study(row)}), 201


@app.route('/api/feasibility-studies/<int:feasibility_id>', methods=['PUT'])
@auth_required
def update_feasibility_study(feasibility_id):
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do estudo de viabilidade e obrigatorio.'}), 400

    conn = get_db()
    existing = get_feasibility_study_for_company(conn, request.current_user['company_id'], feasibility_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Estudo de viabilidade nao encontrado.'}), 404

    conn.execute(
        """
        UPDATE feasibility_studies
        SET title = ?, project_type = ?, sector = ?, sponsor = ?, location = ?, analysis_date = ?,
            horizon_years = ?, capex_estimate = ?, opex_estimate = ?, technical_scope = ?,
            technical_requirements_json = ?, engineering_solution = ?, technology_readiness = ?,
            implementation_schedule = ?, resource_plan_json = ?, regulatory_licensing = ?,
            environmental_social = ?, demand_assumptions = ?, pricing_model = ?,
            revenue_assumptions_json = ?, cost_assumptions_json = ?, financing_structure = ?,
            wacc_assumption = ?, cash_flow_projection = ?, economic_indicators_json = ?,
            sensitivity_analysis = ?, scenario_analysis = ?, key_risks_json = ?,
            risk_response_plan = ?, final_recommendation = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('project_type', ''),
            payload.get('sector', ''),
            payload.get('sponsor', ''),
            payload.get('location', ''),
            payload.get('analysis_date', ''),
            int(payload.get('horizon_years') or 0),
            float(payload.get('capex_estimate') or 0),
            float(payload.get('opex_estimate') or 0),
            payload.get('technical_scope', ''),
            json.dumps(parse_json_list(payload.get('technical_requirements'))),
            payload.get('engineering_solution', ''),
            payload.get('technology_readiness', ''),
            payload.get('implementation_schedule', ''),
            json.dumps(parse_json_list(payload.get('resource_plan'))),
            payload.get('regulatory_licensing', ''),
            payload.get('environmental_social', ''),
            payload.get('demand_assumptions', ''),
            payload.get('pricing_model', ''),
            json.dumps(parse_json_list(payload.get('revenue_assumptions'))),
            json.dumps(parse_json_list(payload.get('cost_assumptions'))),
            payload.get('financing_structure', ''),
            payload.get('wacc_assumption', ''),
            payload.get('cash_flow_projection', ''),
            json.dumps(parse_json_list(payload.get('economic_indicators'))),
            payload.get('sensitivity_analysis', ''),
            payload.get('scenario_analysis', ''),
            json.dumps(parse_json_list(payload.get('key_risks'))),
            payload.get('risk_response_plan', ''),
            payload.get('final_recommendation', ''),
            now_iso(),
            feasibility_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM feasibility_studies WHERE id = ?', (feasibility_id,)).fetchone()
    conn.close()
    return jsonify({'feasibility_study': row_to_feasibility_study(row)})


@app.route('/api/feasibility-studies/<int:feasibility_id>', methods=['DELETE'])
@auth_required
def delete_feasibility_study(feasibility_id):
    conn = get_db()
    existing = get_feasibility_study_for_company(conn, request.current_user['company_id'], feasibility_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Estudo de viabilidade nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM feasibility_studies WHERE id = ? AND company_id = ?',
        (feasibility_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ── Project Management ──────────────────────────────────────────────────────

@app.route('/api/projects', methods=['GET'])
@auth_required
def list_projects():
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM projects WHERE company_id = ? ORDER BY updated_at DESC',
        (request.current_user['company_id'],),
    ).fetchall()
    projects = []
    today = dt.date.today().isoformat()
    for row in rows:
        project = row_to_project(row)
        task_rows = conn.execute(
            'SELECT status, due_date FROM project_tasks WHERE project_id = ? AND company_id = ?',
            (project['id'], request.current_user['company_id']),
        ).fetchall()
        project['task_count'] = len(task_rows)
        project['task_done'] = sum(1 for t in task_rows if t['status'] == 'done')
        project['task_overdue'] = sum(
            1 for t in task_rows
            if t['status'] != 'done' and t['due_date'] and t['due_date'] < today
        )
        projects.append(project)
    conn.close()
    return jsonify({'projects': projects})


@app.route('/api/projects', methods=['POST'])
@auth_required
def create_project():
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do projeto e obrigatorio.'}), 400

    status = (payload.get('status') or 'planning').strip().lower()
    if status not in {'planning', 'active', 'on_hold', 'completed', 'cancelled'}:
        status = 'planning'
    priority = (payload.get('priority') or 'medium').strip().lower()
    if priority not in {'low', 'medium', 'high', 'critical'}:
        priority = 'medium'

    milestones = payload.get('milestones')
    if isinstance(milestones, list):
        clean = []
        for m in milestones:
            if isinstance(m, dict) and (m.get('name') or '').strip():
                clean.append({
                    'name': m['name'].strip(),
                    'due_date': (m.get('due_date') or '').strip(),
                    'status': (m.get('status') or 'pending').strip(),
                })
        milestones = clean
    else:
        milestones = []

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO projects (
            company_id, title, description, status, priority,
            manager_name, manager_email, start_date, end_date, budget,
            objectives_json, milestones_json, risks_json, notes,
            created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            title,
            payload.get('description', ''),
            status,
            priority,
            payload.get('manager_name', ''),
            payload.get('manager_email', ''),
            payload.get('start_date', ''),
            payload.get('end_date', ''),
            float(payload.get('budget') or 0),
            json.dumps(parse_json_list(payload.get('objectives'))),
            json.dumps(milestones, ensure_ascii=False),
            json.dumps(parse_json_list(payload.get('risks'))),
            payload.get('notes', ''),
            timestamp,
            timestamp,
        ),
    )
    project_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM projects WHERE id = ?', (project_id,)).fetchone()
    conn.close()
    return jsonify({'project': row_to_project(row)}), 201


@app.route('/api/projects/<int:project_id>', methods=['PUT'])
@auth_required
def update_project(project_id):
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do projeto e obrigatorio.'}), 400

    conn = get_db()
    existing = get_project_for_company(conn, request.current_user['company_id'], project_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    status = (payload.get('status') or 'planning').strip().lower()
    if status not in {'planning', 'active', 'on_hold', 'completed', 'cancelled'}:
        status = 'planning'
    priority = (payload.get('priority') or 'medium').strip().lower()
    if priority not in {'low', 'medium', 'high', 'critical'}:
        priority = 'medium'

    milestones = payload.get('milestones')
    if isinstance(milestones, list):
        clean = []
        for m in milestones:
            if isinstance(m, dict) and (m.get('name') or '').strip():
                clean.append({
                    'name': m['name'].strip(),
                    'due_date': (m.get('due_date') or '').strip(),
                    'status': (m.get('status') or 'pending').strip(),
                })
        milestones = clean
    else:
        milestones = []

    conn.execute(
        """
        UPDATE projects
        SET title = ?, description = ?, status = ?, priority = ?,
            manager_name = ?, manager_email = ?, start_date = ?, end_date = ?,
            budget = ?, objectives_json = ?, milestones_json = ?, risks_json = ?,
            notes = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('description', ''),
            status,
            priority,
            payload.get('manager_name', ''),
            payload.get('manager_email', ''),
            payload.get('start_date', ''),
            payload.get('end_date', ''),
            float(payload.get('budget') or 0),
            json.dumps(parse_json_list(payload.get('objectives'))),
            json.dumps(milestones, ensure_ascii=False),
            json.dumps(parse_json_list(payload.get('risks'))),
            payload.get('notes', ''),
            now_iso(),
            project_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM projects WHERE id = ?', (project_id,)).fetchone()
    conn.close()
    return jsonify({'project': row_to_project(row)})


@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
@auth_required
def delete_project(project_id):
    conn = get_db()
    existing = get_project_for_company(conn, request.current_user['company_id'], project_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM project_tasks WHERE project_id = ? AND company_id = ?',
        (project_id, request.current_user['company_id']),
    )
    conn.execute(
        'DELETE FROM projects WHERE id = ? AND company_id = ?',
        (project_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ── Project Tasks ────────────────────────────────────────────────────────────

@app.route('/api/projects/<int:project_id>/tasks', methods=['GET'])
@auth_required
def list_project_tasks(project_id):
    conn = get_db()
    existing = get_project_for_company(conn, request.current_user['company_id'], project_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    rows = conn.execute(
        'SELECT * FROM project_tasks WHERE project_id = ? AND company_id = ? ORDER BY due_date ASC, created_at ASC',
        (project_id, request.current_user['company_id']),
    ).fetchall()
    conn.close()
    return jsonify({'tasks': [row_to_task(row) for row in rows]})


@app.route('/api/projects/<int:project_id>/tasks', methods=['POST'])
@auth_required
def create_project_task(project_id):
    conn = get_db()
    existing = get_project_for_company(conn, request.current_user['company_id'], project_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        conn.close()
        return jsonify({'error': 'Titulo da tarefa e obrigatorio.'}), 400

    task_status = (payload.get('status') or 'todo').strip().lower()
    if task_status not in {'todo', 'in_progress', 'review', 'done', 'blocked'}:
        task_status = 'todo'
    task_priority = (payload.get('priority') or 'medium').strip().lower()
    if task_priority not in {'low', 'medium', 'high', 'critical'}:
        task_priority = 'medium'

    completed_date = ''
    if task_status == 'done':
        completed_date = payload.get('completed_date') or now_iso()[:10]

    progress = max(0, min(100, int(payload.get('progress') or 0)))
    if task_status == 'done':
        progress = 100

    timestamp = now_iso()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO project_tasks (
            project_id, company_id, title, description, milestone,
            assignee_name, assignee_email, status, priority,
            start_date, due_date, completed_date, progress,
            created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            project_id,
            request.current_user['company_id'],
            title,
            payload.get('description', ''),
            payload.get('milestone', ''),
            payload.get('assignee_name', ''),
            payload.get('assignee_email', ''),
            task_status,
            task_priority,
            payload.get('start_date', ''),
            payload.get('due_date', ''),
            completed_date,
            progress,
            timestamp,
            timestamp,
        ),
    )
    task_id = cur.lastrowid
    conn.commit()
    row = conn.execute('SELECT * FROM project_tasks WHERE id = ?', (task_id,)).fetchone()
    conn.close()
    return jsonify({'task': row_to_task(row)}), 201


@app.route('/api/projects/<int:project_id>/tasks/<int:task_id>', methods=['PUT'])
@auth_required
def update_project_task(project_id, task_id):
    conn = get_db()
    existing = get_task_for_project(conn, request.current_user['company_id'], project_id, task_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Tarefa nao encontrada.'}), 404

    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        conn.close()
        return jsonify({'error': 'Titulo da tarefa e obrigatorio.'}), 400

    task_status = (payload.get('status') or 'todo').strip().lower()
    if task_status not in {'todo', 'in_progress', 'review', 'done', 'blocked'}:
        task_status = 'todo'
    task_priority = (payload.get('priority') or 'medium').strip().lower()
    if task_priority not in {'low', 'medium', 'high', 'critical'}:
        task_priority = 'medium'

    completed_date = payload.get('completed_date', '')
    if task_status == 'done' and not completed_date:
        completed_date = now_iso()[:10]
    elif task_status != 'done':
        completed_date = ''

    progress = max(0, min(100, int(payload.get('progress') or 0)))
    if task_status == 'done':
        progress = 100

    conn.execute(
        """
        UPDATE project_tasks
        SET title = ?, description = ?, milestone = ?,
            assignee_name = ?, assignee_email = ?, status = ?, priority = ?,
            start_date = ?, due_date = ?, completed_date = ?, progress = ?,
            updated_at = ?
        WHERE id = ? AND project_id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('description', ''),
            payload.get('milestone', ''),
            payload.get('assignee_name', ''),
            payload.get('assignee_email', ''),
            task_status,
            task_priority,
            payload.get('start_date', ''),
            payload.get('due_date', ''),
            completed_date,
            progress,
            now_iso(),
            task_id,
            project_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()
    row = conn.execute('SELECT * FROM project_tasks WHERE id = ?', (task_id,)).fetchone()
    conn.close()
    return jsonify({'task': row_to_task(row)})


@app.route('/api/projects/<int:project_id>/tasks/<int:task_id>', methods=['DELETE'])
@auth_required
def delete_project_task(project_id, task_id):
    conn = get_db()
    existing = get_task_for_project(conn, request.current_user['company_id'], project_id, task_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Tarefa nao encontrada.'}), 404

    conn.execute(
        'DELETE FROM project_tasks WHERE id = ? AND project_id = ? AND company_id = ?',
        (task_id, project_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/plans', methods=['GET'])
@auth_required
def list_plans():
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM strategic_plans WHERE company_id = ? ORDER BY updated_at DESC',
        (request.current_user['company_id'],),
    ).fetchall()
    conn.close()

    plans = [row_to_plan(row) for row in rows]
    return jsonify({'plans': plans})


@app.route('/api/plans', methods=['POST'])
@auth_required
def create_plan():
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do planejamento e obrigatorio.'}), 400

    timestamp = now_iso()
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO strategic_plans (
            company_id, title, period_start, period_end, mission, vision,
            values_json, swot_strengths_json, swot_weaknesses_json, swot_opportunities_json,
            swot_threats_json, objectives_json, initiatives_json, indicators_json,
            created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            request.current_user['company_id'],
            title,
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('mission', ''),
            payload.get('vision', ''),
            json.dumps(parse_json_list(payload.get('values'))),
            json.dumps(parse_json_list(payload.get('swot_strengths'))),
            json.dumps(parse_json_list(payload.get('swot_weaknesses'))),
            json.dumps(parse_json_list(payload.get('swot_opportunities'))),
            json.dumps(parse_json_list(payload.get('swot_threats'))),
            json.dumps(parse_json_list(payload.get('objectives'))),
            json.dumps(parse_json_list(payload.get('initiatives'))),
            json.dumps(parse_json_list(payload.get('indicators'))),
            timestamp,
            timestamp,
        ),
    )
    plan_id = cur.lastrowid
    conn.commit()

    row = conn.execute('SELECT * FROM strategic_plans WHERE id = ?', (plan_id,)).fetchone()
    conn.close()

    return jsonify({'plan': row_to_plan(row)}), 201


@app.route('/api/plans/<int:plan_id>', methods=['PUT'])
@auth_required
def update_plan(plan_id):
    payload = request.get_json(force=True, silent=True) or {}
    title = (payload.get('title') or '').strip()
    if not title:
        return jsonify({'error': 'Titulo do planejamento e obrigatorio.'}), 400

    conn = get_db()
    existing = get_plan_for_company(conn, request.current_user['company_id'], plan_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Planejamento nao encontrado.'}), 404

    conn.execute(
        """
        UPDATE strategic_plans
        SET title = ?, period_start = ?, period_end = ?, mission = ?, vision = ?,
            values_json = ?, swot_strengths_json = ?, swot_weaknesses_json = ?,
            swot_opportunities_json = ?, swot_threats_json = ?, objectives_json = ?,
            initiatives_json = ?, indicators_json = ?, updated_at = ?
        WHERE id = ? AND company_id = ?
        """,
        (
            title,
            payload.get('period_start', ''),
            payload.get('period_end', ''),
            payload.get('mission', ''),
            payload.get('vision', ''),
            json.dumps(parse_json_list(payload.get('values'))),
            json.dumps(parse_json_list(payload.get('swot_strengths'))),
            json.dumps(parse_json_list(payload.get('swot_weaknesses'))),
            json.dumps(parse_json_list(payload.get('swot_opportunities'))),
            json.dumps(parse_json_list(payload.get('swot_threats'))),
            json.dumps(parse_json_list(payload.get('objectives'))),
            json.dumps(parse_json_list(payload.get('initiatives'))),
            json.dumps(parse_json_list(payload.get('indicators'))),
            now_iso(),
            plan_id,
            request.current_user['company_id'],
        ),
    )
    conn.commit()

    row = conn.execute('SELECT * FROM strategic_plans WHERE id = ?', (plan_id,)).fetchone()
    conn.close()

    return jsonify({'plan': row_to_plan(row)})


@app.route('/api/plans/<int:plan_id>', methods=['DELETE'])
@auth_required
def delete_plan(plan_id):
    conn = get_db()
    existing = get_plan_for_company(conn, request.current_user['company_id'], plan_id)
    if not existing:
        conn.close()
        return jsonify({'error': 'Planejamento nao encontrado.'}), 404

    conn.execute(
        'DELETE FROM strategic_plans WHERE id = ? AND company_id = ?',
        (plan_id, request.current_user['company_id']),
    )
    conn.commit()
    conn.close()

    return jsonify({'ok': True})


def load_company_and_plan_or_404(company_id, plan_id):
    conn = get_db()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    plan_row = get_plan_for_company(conn, company_id, plan_id)
    conn.close()

    if not company or not plan_row:
        return None, None

    return row_to_company(company), row_to_plan(plan_row)


def load_company_and_swot_or_404(company_id, swot_id):
    conn = get_db()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    swot_row = get_swot_for_company(conn, company_id, swot_id)
    conn.close()

    if not company or not swot_row:
        return None, None

    return row_to_company(company), row_to_swot(swot_row)


def load_company_and_business_plan_or_404(company_id, business_plan_id):
    conn = get_db()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    business_plan_row = get_business_plan_for_company(conn, company_id, business_plan_id)
    conn.close()

    if not company or not business_plan_row:
        return None, None

    return row_to_company(company), row_to_business_plan(business_plan_row)


def load_company_and_feasibility_or_404(company_id, feasibility_id):
    conn = get_db()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    feasibility_row = get_feasibility_study_for_company(conn, company_id, feasibility_id)
    conn.close()

    if not company or not feasibility_row:
        return None, None

    return row_to_company(company), row_to_feasibility_study(feasibility_row)


def load_company_and_project_or_404(company_id, project_id):
    conn = get_db()
    company = conn.execute('SELECT * FROM companies WHERE id = ?', (company_id,)).fetchone()
    project_row = get_project_for_company(conn, company_id, project_id)
    conn.close()
    if not company or not project_row:
        return None, None
    return row_to_company(company), row_to_project(project_row)


def join_items(items):
    if not items:
        return '-'
    return '\n'.join([f"- {item}" for item in items])


@app.route('/api/feasibility-studies/<int:feasibility_id>/export/docx', methods=['GET'])
@auth_required
def export_feasibility_docx(feasibility_id):
    company, study = load_company_and_feasibility_or_404(request.current_user['company_id'], feasibility_id)
    if not company:
        return jsonify({'error': 'Estudo de viabilidade nao encontrado.'}), 404

    doc = Document()
    doc.add_heading(company.get('name') or 'Empresa', level=1)
    if company.get('header_text'):
        doc.add_paragraph(company['header_text'])

    if company.get('logo_path'):
        logo_file = UPLOADS_DIR / company['logo_path']
        if logo_file.exists():
            doc.add_picture(str(logo_file), width=Inches(1.6))

    doc.add_heading(study['title'], level=2)
    doc.add_paragraph(f"Tipo de projeto: {study.get('project_type') or '-'}")
    doc.add_paragraph(f"Setor: {study.get('sector') or '-'}")
    doc.add_paragraph(f"Patrocinador: {study.get('sponsor') or '-'}")
    doc.add_paragraph(f"Localidade: {study.get('location') or '-'}")
    doc.add_paragraph(f"Data do estudo: {study.get('analysis_date') or '-'}")
    doc.add_paragraph(f"Horizonte (anos): {study.get('horizon_years') or '-'}")
    doc.add_paragraph(f"CAPEX estimado: R$ {study.get('capex_estimate') or 0}")
    doc.add_paragraph(f"OPEX anual estimado: R$ {study.get('opex_estimate') or 0}")

    sections = [
        ('Escopo tecnico', study.get('technical_scope') or '-'),
        ('Requisitos tecnicos criticos', join_items(study.get('technical_requirements') or [])),
        ('Solucao de engenharia', study.get('engineering_solution') or '-'),
        ('Maturidade tecnologica', study.get('technology_readiness') or '-'),
        ('Cronograma de implantacao', study.get('implementation_schedule') or '-'),
        ('Plano de recursos', join_items(study.get('resource_plan') or [])),
        ('Aspectos regulatorios', study.get('regulatory_licensing') or '-'),
        ('Aspectos ambientais e sociais', study.get('environmental_social') or '-'),
        ('Premissas de demanda', study.get('demand_assumptions') or '-'),
        ('Modelo de precificacao', study.get('pricing_model') or '-'),
        ('Premissas de receita', join_items(study.get('revenue_assumptions') or [])),
        ('Premissas de custo', join_items(study.get('cost_assumptions') or [])),
        ('Estrutura de financiamento', study.get('financing_structure') or '-'),
        ('WACC/taxa de desconto', study.get('wacc_assumption') or '-'),
        ('Fluxo de caixa projetado', study.get('cash_flow_projection') or '-'),
        ('Indicadores economico-financeiros', join_items(study.get('economic_indicators') or [])),
        ('Analise de sensibilidade', study.get('sensitivity_analysis') or '-'),
        ('Analise de cenarios', study.get('scenario_analysis') or '-'),
        ('Riscos-chave', join_items(study.get('key_risks') or [])),
        ('Plano de resposta a riscos', study.get('risk_response_plan') or '-'),
        ('Recomendacao final', study.get('final_recommendation') or '-'),
    ]

    for title, content in sections:
        doc.add_heading(title, level=3)
        doc.add_paragraph(content)

    if company.get('footer_text'):
        doc.add_paragraph('')
        doc.add_paragraph(company['footer_text'])

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = secure_filename(f"viabilidade_{feasibility_id}.docx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/feasibility-studies/<int:feasibility_id>/export/pdf', methods=['GET'])
@auth_required
def export_feasibility_pdf(feasibility_id):
    company, study = load_company_and_feasibility_or_404(request.current_user['company_id'], feasibility_id)
    if not company:
        return jsonify({'error': 'Estudo de viabilidade nao encontrado.'}), 404

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def line(text, font='Helvetica', size=10, step=0.58):
        nonlocal y
        if y <= 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text)
        y -= step * cm

    line(company.get('name') or 'Empresa', 'Helvetica-Bold', 14, 0.8)
    if company.get('header_text'):
        line(company['header_text'], 'Helvetica', 10, 0.6)
    line('')

    line(study.get('title') or 'Estudo de Viabilidade', 'Helvetica-Bold', 12, 0.8)
    line(f"Tipo: {study.get('project_type') or '-'}")
    line(f"Setor: {study.get('sector') or '-'}")
    line(f"Patrocinador: {study.get('sponsor') or '-'}")
    line(f"Localidade: {study.get('location') or '-'}")

    blocks = [
        ('Escopo tecnico', [study.get('technical_scope') or '-']),
        ('Requisitos tecnicos', study.get('technical_requirements') or []),
        ('Solucao de engenharia', [study.get('engineering_solution') or '-']),
        ('Maturidade tecnologica', [study.get('technology_readiness') or '-']),
        ('Cronograma de implantacao', [study.get('implementation_schedule') or '-']),
        ('Plano de recursos', study.get('resource_plan') or []),
        ('Aspectos regulatorios', [study.get('regulatory_licensing') or '-']),
        ('Aspectos ambientais e sociais', [study.get('environmental_social') or '-']),
        ('Premissas de demanda', [study.get('demand_assumptions') or '-']),
        ('Modelo de precificacao', [study.get('pricing_model') or '-']),
        ('Premissas de receita', study.get('revenue_assumptions') or []),
        ('Premissas de custo', study.get('cost_assumptions') or []),
        ('Estrutura de financiamento', [study.get('financing_structure') or '-']),
        ('WACC/taxa de desconto', [study.get('wacc_assumption') or '-']),
        ('Fluxo de caixa projetado', [study.get('cash_flow_projection') or '-']),
        ('Indicadores economico-financeiros', study.get('economic_indicators') or []),
        ('Analise de sensibilidade', [study.get('sensitivity_analysis') or '-']),
        ('Analise de cenarios', [study.get('scenario_analysis') or '-']),
        ('Riscos-chave', study.get('key_risks') or []),
        ('Plano de resposta a riscos', [study.get('risk_response_plan') or '-']),
        ('Recomendacao final', [study.get('final_recommendation') or '-']),
    ]

    for title, items in blocks:
        line('')
        line(title, 'Helvetica-Bold', 11, 0.7)
        if not items:
            line('-')
            continue
        for item in items:
            txt = f"- {item}"
            if len(txt) <= 95:
                line(txt)
                continue
            while txt:
                chunk = txt[:95]
                txt = txt[95:]
                line(chunk)

    if company.get('footer_text'):
        line('')
        line(company['footer_text'], 'Helvetica-Oblique', 9)

    pdf.save()
    stream.seek(0)

    filename = secure_filename(f"viabilidade_{feasibility_id}.pdf")
    return send_file(stream, as_attachment=True, download_name=filename, mimetype='application/pdf')


@app.route('/api/feasibility-studies/<int:feasibility_id>/export/excel', methods=['GET'])
@auth_required
def export_feasibility_excel(feasibility_id):
    company, study = load_company_and_feasibility_or_404(request.current_user['company_id'], feasibility_id)
    if not company:
        return jsonify({'error': 'Estudo de viabilidade nao encontrado.'}), 404

    wb = Workbook()
    ws = wb.active
    ws.title = 'Viabilidade'

    ws['A1'] = company.get('name') or ''
    ws['A2'] = company.get('header_text') or ''
    ws['A4'] = study.get('title') or 'Estudo de Viabilidade'
    ws['A5'] = f"Tipo: {study.get('project_type') or '-'}"
    ws['A6'] = f"Setor: {study.get('sector') or '-'}"
    ws['A7'] = f"Patrocinador: {study.get('sponsor') or '-'}"
    ws['A8'] = f"Localidade: {study.get('location') or '-'}"
    ws['A9'] = f"Data: {study.get('analysis_date') or '-'}"
    ws['A10'] = f"Horizonte (anos): {study.get('horizon_years') or '-'}"
    ws['A11'] = f"CAPEX: R$ {study.get('capex_estimate') or 0}"
    ws['A12'] = f"OPEX: R$ {study.get('opex_estimate') or 0}"

    rows = [
        ('Escopo tecnico', study.get('technical_scope') or '-'),
        ('Requisitos tecnicos', join_items(study.get('technical_requirements') or [])),
        ('Solucao de engenharia', study.get('engineering_solution') or '-'),
        ('Maturidade tecnologica', study.get('technology_readiness') or '-'),
        ('Cronograma de implantacao', study.get('implementation_schedule') or '-'),
        ('Plano de recursos', join_items(study.get('resource_plan') or [])),
        ('Aspectos regulatorios', study.get('regulatory_licensing') or '-'),
        ('Aspectos ambientais e sociais', study.get('environmental_social') or '-'),
        ('Premissas de demanda', study.get('demand_assumptions') or '-'),
        ('Modelo de precificacao', study.get('pricing_model') or '-'),
        ('Premissas de receita', join_items(study.get('revenue_assumptions') or [])),
        ('Premissas de custo', join_items(study.get('cost_assumptions') or [])),
        ('Estrutura de financiamento', study.get('financing_structure') or '-'),
        ('WACC/taxa de desconto', study.get('wacc_assumption') or '-'),
        ('Fluxo de caixa projetado', study.get('cash_flow_projection') or '-'),
        ('Indicadores economico-financeiros', join_items(study.get('economic_indicators') or [])),
        ('Analise de sensibilidade', study.get('sensitivity_analysis') or '-'),
        ('Analise de cenarios', study.get('scenario_analysis') or '-'),
        ('Riscos-chave', join_items(study.get('key_risks') or [])),
        ('Plano de resposta a riscos', study.get('risk_response_plan') or '-'),
        ('Recomendacao final', study.get('final_recommendation') or '-'),
    ]

    start_row = 14
    for idx, (section, content) in enumerate(rows):
        row_number = start_row + idx
        ws[f'A{row_number}'] = section
        ws[f'B{row_number}'] = content

    ws['A40'] = company.get('footer_text') or ''

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = secure_filename(f"viabilidade_{feasibility_id}.xlsx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


@app.route('/api/business-plans/<int:business_plan_id>/export/docx', methods=['GET'])
@auth_required
def export_business_plan_docx(business_plan_id):
    company, business_plan = load_company_and_business_plan_or_404(request.current_user['company_id'], business_plan_id)
    if not company:
        return jsonify({'error': 'Plano de negocios nao encontrado.'}), 404

    doc = Document()
    doc.add_heading(company.get('name') or 'Empresa', level=1)
    if company.get('header_text'):
        doc.add_paragraph(company['header_text'])

    if company.get('logo_path'):
        logo_file = UPLOADS_DIR / company['logo_path']
        if logo_file.exists():
            doc.add_picture(str(logo_file), width=Inches(1.6))

    doc.add_heading(business_plan['title'], level=2)
    doc.add_paragraph(f"Segmento: {business_plan.get('industry') or '-'}")
    doc.add_paragraph(f"Publico-alvo: {business_plan.get('target_market') or '-'}")
    doc.add_paragraph(f"Periodo: {business_plan.get('period_start') or '-'} ate {business_plan.get('period_end') or '-'}")

    sections = [
        ('Resumo executivo', business_plan.get('executive_summary') or '-'),
        ('Problema que o negocio resolve', business_plan.get('problem_statement') or '-'),
        ('Proposta de valor', business_plan.get('value_proposition') or '-'),
        ('Produtos e servicos', join_items(business_plan.get('products_services') or [])),
        ('Analise de mercado e concorrencia', business_plan.get('market_analysis') or '-'),
        ('Estrategia comercial e marketing', business_plan.get('marketing_strategy') or '-'),
        ('Plano operacional', business_plan.get('operational_plan') or '-'),
        ('Modelo de receita', join_items(business_plan.get('revenue_streams') or [])),
        ('Estrutura de custos', join_items(business_plan.get('cost_structure') or [])),
        ('Necessidade de investimento', join_items(business_plan.get('investment_requirements') or [])),
        ('Projecoes financeiras', join_items(business_plan.get('financial_projections') or [])),
        ('Estrutura de equipe', join_items(business_plan.get('team_structure') or [])),
        ('Marcos estrategicos', join_items(business_plan.get('milestones') or [])),
        ('Riscos e mitigacoes', join_items(business_plan.get('risks_mitigation') or [])),
    ]

    for title, content in sections:
        doc.add_heading(title, level=3)
        doc.add_paragraph(content)

    if company.get('footer_text'):
        doc.add_paragraph('')
        doc.add_paragraph(company['footer_text'])

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = secure_filename(f"plano_negocios_{business_plan_id}.docx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/business-plans/<int:business_plan_id>/export/pdf', methods=['GET'])
@auth_required
def export_business_plan_pdf(business_plan_id):
    company, business_plan = load_company_and_business_plan_or_404(request.current_user['company_id'], business_plan_id)
    if not company:
        return jsonify({'error': 'Plano de negocios nao encontrado.'}), 404

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def line(text, font='Helvetica', size=10, step=0.58):
        nonlocal y
        if y <= 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text)
        y -= step * cm

    line(company.get('name') or 'Empresa', 'Helvetica-Bold', 14, 0.8)
    if company.get('header_text'):
        line(company['header_text'], 'Helvetica', 10, 0.6)
    line('')

    line(business_plan.get('title') or 'Plano de Negocios', 'Helvetica-Bold', 12, 0.8)
    line(f"Segmento: {business_plan.get('industry') or '-'}")
    line(f"Publico-alvo: {business_plan.get('target_market') or '-'}")
    line(f"Periodo: {business_plan.get('period_start') or '-'} ate {business_plan.get('period_end') or '-'}")

    blocks = [
        ('Resumo executivo', [business_plan.get('executive_summary') or '-']),
        ('Problema que o negocio resolve', [business_plan.get('problem_statement') or '-']),
        ('Proposta de valor', [business_plan.get('value_proposition') or '-']),
        ('Produtos e servicos', business_plan.get('products_services') or []),
        ('Analise de mercado e concorrencia', [business_plan.get('market_analysis') or '-']),
        ('Estrategia comercial e marketing', [business_plan.get('marketing_strategy') or '-']),
        ('Plano operacional', [business_plan.get('operational_plan') or '-']),
        ('Modelo de receita', business_plan.get('revenue_streams') or []),
        ('Estrutura de custos', business_plan.get('cost_structure') or []),
        ('Necessidade de investimento', business_plan.get('investment_requirements') or []),
        ('Projecoes financeiras', business_plan.get('financial_projections') or []),
        ('Estrutura de equipe', business_plan.get('team_structure') or []),
        ('Marcos estrategicos', business_plan.get('milestones') or []),
        ('Riscos e mitigacoes', business_plan.get('risks_mitigation') or []),
    ]

    for title, items in blocks:
        line('')
        line(title, 'Helvetica-Bold', 11, 0.7)
        if not items:
            line('-')
            continue
        for item in items:
            txt = f"- {item}"
            if len(txt) <= 95:
                line(txt)
                continue
            while txt:
                chunk = txt[:95]
                txt = txt[95:]
                line(chunk)

    if company.get('footer_text'):
        line('')
        line(company['footer_text'], 'Helvetica-Oblique', 9)

    pdf.save()
    stream.seek(0)

    filename = secure_filename(f"plano_negocios_{business_plan_id}.pdf")
    return send_file(stream, as_attachment=True, download_name=filename, mimetype='application/pdf')


@app.route('/api/business-plans/<int:business_plan_id>/export/excel', methods=['GET'])
@auth_required
def export_business_plan_excel(business_plan_id):
    company, business_plan = load_company_and_business_plan_or_404(request.current_user['company_id'], business_plan_id)
    if not company:
        return jsonify({'error': 'Plano de negocios nao encontrado.'}), 404

    wb = Workbook()
    ws = wb.active
    ws.title = 'Plano de Negocios'

    ws['A1'] = company.get('name') or ''
    ws['A2'] = company.get('header_text') or ''
    ws['A4'] = business_plan.get('title') or 'Plano de Negocios'
    ws['A5'] = f"Segmento: {business_plan.get('industry') or '-'}"
    ws['A6'] = f"Publico-alvo: {business_plan.get('target_market') or '-'}"
    ws['A7'] = f"Periodo: {business_plan.get('period_start') or '-'} ate {business_plan.get('period_end') or '-'}"

    rows = [
        ('Resumo executivo', business_plan.get('executive_summary') or '-'),
        ('Problema que o negocio resolve', business_plan.get('problem_statement') or '-'),
        ('Proposta de valor', business_plan.get('value_proposition') or '-'),
        ('Produtos e servicos', join_items(business_plan.get('products_services') or [])),
        ('Analise de mercado e concorrencia', business_plan.get('market_analysis') or '-'),
        ('Estrategia comercial e marketing', business_plan.get('marketing_strategy') or '-'),
        ('Plano operacional', business_plan.get('operational_plan') or '-'),
        ('Modelo de receita', join_items(business_plan.get('revenue_streams') or [])),
        ('Estrutura de custos', join_items(business_plan.get('cost_structure') or [])),
        ('Necessidade de investimento', join_items(business_plan.get('investment_requirements') or [])),
        ('Projecoes financeiras', join_items(business_plan.get('financial_projections') or [])),
        ('Estrutura de equipe', join_items(business_plan.get('team_structure') or [])),
        ('Marcos estrategicos', join_items(business_plan.get('milestones') or [])),
        ('Riscos e mitigacoes', join_items(business_plan.get('risks_mitigation') or [])),
    ]

    start_row = 9
    for idx, (section, content) in enumerate(rows):
        row_number = start_row + idx
        ws[f'A{row_number}'] = section
        ws[f'B{row_number}'] = content

    ws['A28'] = company.get('footer_text') or ''

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = secure_filename(f"plano_negocios_{business_plan_id}.xlsx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


@app.route('/api/swot-analyses/<int:swot_id>/export/docx', methods=['GET'])
@auth_required
def export_swot_docx(swot_id):
    company, swot = load_company_and_swot_or_404(request.current_user['company_id'], swot_id)
    if not company:
        return jsonify({'error': 'Analise SWOT nao encontrada.'}), 404

    doc = Document()
    doc.add_heading(company.get('name') or 'Empresa', level=1)
    if company.get('header_text'):
        doc.add_paragraph(company['header_text'])

    if company.get('logo_path'):
        logo_file = UPLOADS_DIR / company['logo_path']
        if logo_file.exists():
            doc.add_picture(str(logo_file), width=Inches(1.6))

    doc.add_heading(swot['title'], level=2)
    doc.add_paragraph(f"Data da analise: {swot.get('analysis_date') or '-'}")
    doc.add_paragraph(f"Area: {swot.get('business_unit') or '-'}")
    doc.add_paragraph(f"Periodo: {swot.get('period_start') or '-'} ate {swot.get('period_end') or '-'}")
    doc.add_paragraph(f"Objetivo: {swot.get('analysis_objective') or '-'}")

    sections = [
        ('Forcas', join_items(swot.get('strengths') or [])),
        ('Fraquezas', join_items(swot.get('weaknesses') or [])),
        ('Oportunidades', join_items(swot.get('opportunities') or [])),
        ('Ameacas', join_items(swot.get('threats') or [])),
        ('Estrategias FO', join_items(swot.get('so_strategies') or [])),
        ('Estrategias DO', join_items(swot.get('wo_strategies') or [])),
        ('Estrategias FA', join_items(swot.get('st_strategies') or [])),
        ('Estrategias DA', join_items(swot.get('wt_strategies') or [])),
        ('Plano tatico prioritario', join_items(swot.get('priority_actions') or [])),
        ('Riscos criticos e mitigacoes', join_items(swot.get('critical_risks') or [])),
        ('Parecer executivo', swot.get('executive_summary') or '-'),
    ]

    for title, content in sections:
        doc.add_heading(title, level=3)
        doc.add_paragraph(content)

    if company.get('footer_text'):
        doc.add_paragraph('')
        doc.add_paragraph(company['footer_text'])

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = secure_filename(f"swot_{swot_id}.docx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/swot-analyses/<int:swot_id>/export/pdf', methods=['GET'])
@auth_required
def export_swot_pdf(swot_id):
    company, swot = load_company_and_swot_or_404(request.current_user['company_id'], swot_id)
    if not company:
        return jsonify({'error': 'Analise SWOT nao encontrada.'}), 404

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def line(text, font='Helvetica', size=10, step=0.58):
        nonlocal y
        if y <= 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text)
        y -= step * cm

    line(company.get('name') or 'Empresa', 'Helvetica-Bold', 14, 0.8)
    if company.get('header_text'):
        line(company['header_text'], 'Helvetica', 10, 0.6)
    line('')

    line(swot.get('title') or 'Analise SWOT', 'Helvetica-Bold', 12, 0.8)
    line(f"Data: {swot.get('analysis_date') or '-'}")
    line(f"Area: {swot.get('business_unit') or '-'}")
    line(f"Periodo: {swot.get('period_start') or '-'} ate {swot.get('period_end') or '-'}")
    line(f"Objetivo: {swot.get('analysis_objective') or '-'}")

    blocks = [
        ('Forcas', swot.get('strengths') or []),
        ('Fraquezas', swot.get('weaknesses') or []),
        ('Oportunidades', swot.get('opportunities') or []),
        ('Ameacas', swot.get('threats') or []),
        ('Estrategias FO', swot.get('so_strategies') or []),
        ('Estrategias DO', swot.get('wo_strategies') or []),
        ('Estrategias FA', swot.get('st_strategies') or []),
        ('Estrategias DA', swot.get('wt_strategies') or []),
        ('Plano tatico prioritario', swot.get('priority_actions') or []),
        ('Riscos criticos e mitigacoes', swot.get('critical_risks') or []),
        ('Parecer executivo', [swot.get('executive_summary') or '-']),
    ]

    for title, items in blocks:
        line('')
        line(title, 'Helvetica-Bold', 11, 0.7)
        if not items:
            line('-')
            continue
        for item in items:
            txt = f"- {item}"
            if len(txt) <= 95:
                line(txt)
                continue
            while txt:
                chunk = txt[:95]
                txt = txt[95:]
                line(chunk)

    if company.get('footer_text'):
        line('')
        line(company['footer_text'], 'Helvetica-Oblique', 9)

    pdf.save()
    stream.seek(0)

    filename = secure_filename(f"swot_{swot_id}.pdf")
    return send_file(stream, as_attachment=True, download_name=filename, mimetype='application/pdf')


@app.route('/api/swot-analyses/<int:swot_id>/export/excel', methods=['GET'])
@auth_required
def export_swot_excel(swot_id):
    company, swot = load_company_and_swot_or_404(request.current_user['company_id'], swot_id)
    if not company:
        return jsonify({'error': 'Analise SWOT nao encontrada.'}), 404

    wb = Workbook()
    ws = wb.active
    ws.title = 'SWOT'

    ws['A1'] = company.get('name') or ''
    ws['A2'] = company.get('header_text') or ''
    ws['A4'] = swot.get('title') or 'Analise SWOT'
    ws['A5'] = f"Data: {swot.get('analysis_date') or '-'}"
    ws['A6'] = f"Area: {swot.get('business_unit') or '-'}"
    ws['A7'] = f"Periodo: {swot.get('period_start') or '-'} ate {swot.get('period_end') or '-'}"
    ws['A8'] = f"Objetivo: {swot.get('analysis_objective') or '-'}"

    rows = [
        ('Forcas', join_items(swot.get('strengths') or [])),
        ('Fraquezas', join_items(swot.get('weaknesses') or [])),
        ('Oportunidades', join_items(swot.get('opportunities') or [])),
        ('Ameacas', join_items(swot.get('threats') or [])),
        ('Estrategias FO', join_items(swot.get('so_strategies') or [])),
        ('Estrategias DO', join_items(swot.get('wo_strategies') or [])),
        ('Estrategias FA', join_items(swot.get('st_strategies') or [])),
        ('Estrategias DA', join_items(swot.get('wt_strategies') or [])),
        ('Plano tatico prioritario', join_items(swot.get('priority_actions') or [])),
        ('Riscos criticos e mitigacoes', join_items(swot.get('critical_risks') or [])),
        ('Parecer executivo', swot.get('executive_summary') or '-'),
    ]

    start_row = 10
    for idx, (section, content) in enumerate(rows):
        row_number = start_row + idx
        ws[f'A{row_number}'] = section
        ws[f'B{row_number}'] = content

    ws['A30'] = company.get('footer_text') or ''

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = secure_filename(f"swot_{swot_id}.xlsx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )


@app.route('/api/plans/<int:plan_id>/export/docx', methods=['GET'])
@auth_required
def export_docx(plan_id):
    company, plan = load_company_and_plan_or_404(request.current_user['company_id'], plan_id)
    if not company:
        return jsonify({'error': 'Planejamento nao encontrado.'}), 404

    doc = Document()

    doc.add_heading(company.get('name') or 'Empresa', level=1)
    if company.get('header_text'):
        doc.add_paragraph(company['header_text'])

    if company.get('logo_path'):
        logo_file = UPLOADS_DIR / company['logo_path']
        if logo_file.exists():
            doc.add_picture(str(logo_file), width=Inches(1.6))

    doc.add_heading(plan['title'], level=2)
    period = f"Periodo: {plan.get('period_start') or '-'} ate {plan.get('period_end') or '-'}"
    doc.add_paragraph(period)

    sections = [
        ('Missao', plan.get('mission') or '-'),
        ('Visao', plan.get('vision') or '-'),
        ('Valores', join_items(plan.get('values') or [])),
        ('SWOT - Forcas', join_items(plan.get('swot_strengths') or [])),
        ('SWOT - Fraquezas', join_items(plan.get('swot_weaknesses') or [])),
        ('SWOT - Oportunidades', join_items(plan.get('swot_opportunities') or [])),
        ('SWOT - Ameacas', join_items(plan.get('swot_threats') or [])),
        ('Objetivos Estrategicos', join_items(plan.get('objectives') or [])),
        ('Iniciativas', join_items(plan.get('initiatives') or [])),
        ('Indicadores (KPIs)', join_items(plan.get('indicators') or [])),
    ]

    for title, content in sections:
        doc.add_heading(title, level=3)
        doc.add_paragraph(content)

    if company.get('footer_text'):
        doc.add_paragraph('')
        doc.add_paragraph(company['footer_text'])

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = secure_filename(f"planejamento_{plan_id}.docx")
    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/plans/<int:plan_id>/export/pdf', methods=['GET'])
@auth_required
def export_pdf(plan_id):
    company, plan = load_company_and_plan_or_404(request.current_user['company_id'], plan_id)
    if not company:
        return jsonify({'error': 'Planejamento nao encontrado.'}), 404

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def line(text, font='Helvetica', size=10, step=0.58):
        nonlocal y
        if y <= 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text)
        y -= step * cm

    line(company.get('name') or 'Empresa', 'Helvetica-Bold', 14, 0.8)
    if company.get('header_text'):
        line(company['header_text'], 'Helvetica', 10, 0.6)
    line('')

    line(plan.get('title') or 'Planejamento Estrategico', 'Helvetica-Bold', 12, 0.8)
    line(f"Periodo: {plan.get('period_start') or '-'} ate {plan.get('period_end') or '-'}")

    blocks = [
        ('Missao', [plan.get('mission') or '-']),
        ('Visao', [plan.get('vision') or '-']),
        ('Valores', plan.get('values') or []),
        ('SWOT - Forcas', plan.get('swot_strengths') or []),
        ('SWOT - Fraquezas', plan.get('swot_weaknesses') or []),
        ('SWOT - Oportunidades', plan.get('swot_opportunities') or []),
        ('SWOT - Ameacas', plan.get('swot_threats') or []),
        ('Objetivos', plan.get('objectives') or []),
        ('Iniciativas', plan.get('initiatives') or []),
        ('Indicadores', plan.get('indicators') or []),
    ]

    for title, items in blocks:
        line('')
        line(title, 'Helvetica-Bold', 11, 0.7)
        if not items:
            line('-')
            continue
        for item in items:
            txt = f"- {item}"
            if len(txt) <= 95:
                line(txt)
                continue
            # quebra simples para manter legibilidade sem dependencias extras
            while txt:
                chunk = txt[:95]
                txt = txt[95:]
                line(chunk)

    if company.get('footer_text'):
        line('')
        line(company['footer_text'], 'Helvetica-Oblique', 9)

    pdf.save()
    stream.seek(0)

    filename = secure_filename(f"planejamento_{plan_id}.pdf")
    return send_file(stream, as_attachment=True, download_name=filename, mimetype='application/pdf')


@app.route('/api/plans/<int:plan_id>/export/excel', methods=['GET'])
@auth_required
def export_excel(plan_id):
    company, plan = load_company_and_plan_or_404(request.current_user['company_id'], plan_id)
    if not company:
        return jsonify({'error': 'Planejamento nao encontrado.'}), 404

    use_template = Path(TEMPLATE_PATH).exists()

    if use_template:
        wb = load_workbook(TEMPLATE_PATH, keep_vba=True)
        ws = wb[wb.sheetnames[0]]
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = 'Planejamento'

    ws['A1'] = company.get('name') or ''
    ws['A2'] = company.get('header_text') or ''
    ws['A3'] = f"CNPJ: {company.get('cnpj') or '-'}"
    ws['A5'] = plan.get('title') or ''
    ws['A6'] = f"Periodo: {plan.get('period_start') or '-'} ate {plan.get('period_end') or '-'}"

    rows = [
        ('Missao', plan.get('mission') or '-'),
        ('Visao', plan.get('vision') or '-'),
        ('Valores', join_items(plan.get('values') or [])),
        ('SWOT Forcas', join_items(plan.get('swot_strengths') or [])),
        ('SWOT Fraquezas', join_items(plan.get('swot_weaknesses') or [])),
        ('SWOT Oportunidades', join_items(plan.get('swot_opportunities') or [])),
        ('SWOT Ameacas', join_items(plan.get('swot_threats') or [])),
        ('Objetivos', join_items(plan.get('objectives') or [])),
        ('Iniciativas', join_items(plan.get('initiatives') or [])),
        ('Indicadores', join_items(plan.get('indicators') or [])),
    ]

    start_row = 8
    for idx, (section, content) in enumerate(rows):
        row_number = start_row + idx
        ws[f'A{row_number}'] = section
        ws[f'B{row_number}'] = content

    ws['A24'] = company.get('footer_text') or ''

    stream = io.BytesIO()
    if use_template:
        wb.save(stream)
        ext = 'xlsm'
        mime = 'application/vnd.ms-excel.sheet.macroEnabled.12'
    else:
        wb.save(stream)
        ext = 'xlsx'
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

    stream.seek(0)
    filename = secure_filename(f"planejamento_{plan_id}.{ext}")

    return send_file(stream, as_attachment=True, download_name=filename, mimetype=mime)


# ── Project Export Endpoints ─────────────────────────────────────────────────

def _get_project_with_tasks(company_id, project_id):
    company, project = load_company_and_project_or_404(company_id, project_id)
    if not company:
        return None, None, []
    conn = get_db()
    task_rows = conn.execute(
        'SELECT * FROM project_tasks WHERE project_id = ? AND company_id = ? ORDER BY due_date ASC',
        (project_id, company_id),
    ).fetchall()
    conn.close()
    tasks = [row_to_task(row) for row in task_rows]
    return company, project, tasks


def _status_label(val):
    labels = {
        'planning': 'Planejamento', 'active': 'Ativo', 'on_hold': 'Em espera',
        'completed': 'Concluido', 'cancelled': 'Cancelado',
        'todo': 'A Fazer', 'in_progress': 'Em Andamento', 'review': 'Em Revisao',
        'done': 'Concluida', 'blocked': 'Bloqueada',
    }
    return labels.get(val, val or '-')


def _priority_label(val):
    labels = {'low': 'Baixa', 'medium': 'Media', 'high': 'Alta', 'critical': 'Critica'}
    return labels.get(val, val or '-')


@app.route('/api/projects/<int:project_id>/export/docx', methods=['GET'])
@auth_required
def export_project_docx(project_id):
    company, project, tasks = _get_project_with_tasks(request.current_user['company_id'], project_id)
    if not company:
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    doc = Document()
    doc.add_heading(company.get('name') or 'Empresa', level=1)
    if company.get('header_text'):
        doc.add_paragraph(company['header_text'])
    if company.get('logo_path'):
        logo_file = UPLOADS_DIR / company['logo_path']
        if logo_file.exists():
            doc.add_picture(str(logo_file), width=Inches(1.6))

    doc.add_heading(f"Projeto: {project['title']}", level=2)
    doc.add_paragraph(f"Status: {_status_label(project.get('status'))}")
    doc.add_paragraph(f"Prioridade: {_priority_label(project.get('priority'))}")
    doc.add_paragraph(f"Responsavel: {project.get('manager_name') or '-'} ({project.get('manager_email') or '-'})")
    doc.add_paragraph(f"Periodo: {project.get('start_date') or '-'} ate {project.get('end_date') or '-'}")
    doc.add_paragraph(f"Orcamento: R$ {project.get('budget') or 0}")

    if project.get('description'):
        doc.add_heading('Descricao', level=3)
        doc.add_paragraph(project['description'])

    if project.get('objectives'):
        doc.add_heading('Objetivos e Metas', level=3)
        doc.add_paragraph(join_items(project['objectives']))

    milestones = project.get('milestones') or []
    if milestones:
        doc.add_heading('Etapas / Marcos', level=3)
        for m in milestones:
            doc.add_paragraph(f"- {m.get('name', '')} | Prazo: {m.get('due_date') or '-'} | Status: {_status_label(m.get('status'))}")

    if project.get('risks'):
        doc.add_heading('Riscos Identificados', level=3)
        doc.add_paragraph(join_items(project['risks']))

    if project.get('notes'):
        doc.add_heading('Observacoes', level=3)
        doc.add_paragraph(project['notes'])

    if tasks:
        doc.add_heading('Tarefas do Projeto', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ['Tarefa', 'Responsavel', 'Etapa', 'Status', 'Prazo', 'Progresso']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for t in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = t.get('title') or ''
            row_cells[1].text = t.get('assignee_name') or '-'
            row_cells[2].text = t.get('milestone') or '-'
            row_cells[3].text = _status_label(t.get('status'))
            row_cells[4].text = t.get('due_date') or '-'
            row_cells[5].text = f"{t.get('progress', 0)}%"

    if company.get('footer_text'):
        doc.add_paragraph('')
        doc.add_paragraph(company['footer_text'])

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    filename = secure_filename(f"projeto_{project_id}.docx")
    return send_file(stream, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/projects/<int:project_id>/export/pdf', methods=['GET'])
@auth_required
def export_project_pdf(project_id):
    company, project, tasks = _get_project_with_tasks(request.current_user['company_id'], project_id)
    if not company:
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def line(text, font='Helvetica', size=10, step=0.58):
        nonlocal y
        if y <= 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text)
        y -= step * cm

    line(company.get('name') or 'Empresa', 'Helvetica-Bold', 14, 0.8)
    if company.get('header_text'):
        line(company['header_text'], 'Helvetica', 10, 0.6)
    line('')
    line(f"Projeto: {project.get('title', '')}", 'Helvetica-Bold', 12, 0.8)
    line(f"Status: {_status_label(project.get('status'))} | Prioridade: {_priority_label(project.get('priority'))}")
    line(f"Responsavel: {project.get('manager_name') or '-'}")
    line(f"Periodo: {project.get('start_date') or '-'} ate {project.get('end_date') or '-'}")
    line(f"Orcamento: R$ {project.get('budget') or 0}")

    if project.get('description'):
        line('')
        line('Descricao', 'Helvetica-Bold', 11, 0.7)
        for chunk_line in (project['description'] or '').split('\n'):
            txt = chunk_line.strip()
            while txt:
                line(txt[:95])
                txt = txt[95:]

    objectives = project.get('objectives') or []
    if objectives:
        line('')
        line('Objetivos e Metas', 'Helvetica-Bold', 11, 0.7)
        for obj in objectives:
            line(f"- {obj}"[:95])

    milestones = project.get('milestones') or []
    if milestones:
        line('')
        line('Etapas / Marcos', 'Helvetica-Bold', 11, 0.7)
        for m in milestones:
            line(f"- {m.get('name', '')} | {m.get('due_date') or '-'} | {_status_label(m.get('status'))}"[:95])

    if tasks:
        line('')
        line('Tarefas do Projeto', 'Helvetica-Bold', 11, 0.7)
        for t in tasks:
            line(f"- {t.get('title', '')} | {t.get('assignee_name') or '-'} | {_status_label(t.get('status'))} | {t.get('due_date') or '-'} | {t.get('progress', 0)}%"[:95])

    if company.get('footer_text'):
        line('')
        line(company['footer_text'], 'Helvetica-Oblique', 9)

    pdf.save()
    stream.seek(0)
    filename = secure_filename(f"projeto_{project_id}.pdf")
    return send_file(stream, as_attachment=True, download_name=filename, mimetype='application/pdf')


@app.route('/api/projects/<int:project_id>/export/excel', methods=['GET'])
@auth_required
def export_project_excel(project_id):
    company, project, tasks = _get_project_with_tasks(request.current_user['company_id'], project_id)
    if not company:
        return jsonify({'error': 'Projeto nao encontrado.'}), 404

    wb = Workbook()
    ws = wb.active
    ws.title = 'Projeto'

    ws['A1'] = company.get('name') or ''
    ws['A2'] = company.get('header_text') or ''
    ws['A4'] = f"Projeto: {project.get('title', '')}"
    ws['A5'] = f"Status: {_status_label(project.get('status'))}"
    ws['A6'] = f"Prioridade: {_priority_label(project.get('priority'))}"
    ws['A7'] = f"Responsavel: {project.get('manager_name') or '-'}"
    ws['A8'] = f"Periodo: {project.get('start_date') or '-'} ate {project.get('end_date') or '-'}"
    ws['A9'] = f"Orcamento: R$ {project.get('budget') or 0}"
    ws['A10'] = f"Descricao: {project.get('description') or '-'}"
    ws['A11'] = f"Objetivos: {join_items(project.get('objectives') or [])}"

    milestones = project.get('milestones') or []
    r = 13
    ws[f'A{r}'] = 'ETAPAS / MARCOS'
    r += 1
    ws[f'A{r}'] = 'Nome'
    ws[f'B{r}'] = 'Prazo'
    ws[f'C{r}'] = 'Status'
    r += 1
    for m in milestones:
        ws[f'A{r}'] = m.get('name', '')
        ws[f'B{r}'] = m.get('due_date', '')
        ws[f'C{r}'] = _status_label(m.get('status'))
        r += 1

    r += 1
    ws[f'A{r}'] = 'TAREFAS DO PROJETO'
    r += 1
    headers = ['Tarefa', 'Responsavel', 'Etapa', 'Status', 'Prioridade', 'Inicio', 'Prazo', 'Conclusao', 'Progresso']
    for i, h in enumerate(headers):
        ws.cell(row=r, column=i + 1, value=h)
    r += 1
    for t in tasks:
        ws.cell(row=r, column=1, value=t.get('title', ''))
        ws.cell(row=r, column=2, value=t.get('assignee_name', ''))
        ws.cell(row=r, column=3, value=t.get('milestone', ''))
        ws.cell(row=r, column=4, value=_status_label(t.get('status')))
        ws.cell(row=r, column=5, value=_priority_label(t.get('priority')))
        ws.cell(row=r, column=6, value=t.get('start_date', ''))
        ws.cell(row=r, column=7, value=t.get('due_date', ''))
        ws.cell(row=r, column=8, value=t.get('completed_date', ''))
        ws.cell(row=r, column=9, value=f"{t.get('progress', 0)}%")
        r += 1

    ws[f'A{r + 1}'] = company.get('footer_text') or ''

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    filename = secure_filename(f"projeto_{project_id}.xlsx")
    return send_file(stream, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


init_db()
migrate_db()
migrate_scheduled_status()

# ══════════════════════════════════════════════════════════
# NEW ENDPOINTS — CMS v2
# ══════════════════════════════════════════════════════════

# ── Single content by ID ──────────────────────────────────
@app.route('/api/content/<int:content_id>', methods=['GET'])
@auth_required
def get_content_by_id(content_id):
    conn = get_db()
    row = get_content_for_company(conn, request.current_user['company_id'], content_id)
    conn.close()
    if not row:
        return jsonify({'error': 'Conteudo nao encontrado.'}), 404
    return jsonify({'item': row_to_content(row)})


# ── Change password ───────────────────────────────────────
@app.route('/api/auth/change-password', methods=['POST'])
@auth_required
def change_password():
    payload = request.get_json(force=True, silent=True) or {}
    old_pw = payload.get('old_password') or ''
    new_pw = payload.get('new_password') or ''
    if not old_pw or len(new_pw) < 8:
        return jsonify({'error': 'Senha atual obrigatoria e nova senha deve ter ao menos 8 caracteres.'}), 400
    conn = get_db()
    user = conn.execute('SELECT id, password_hash FROM users WHERE id = ?',
                        (request.current_user['id'],)).fetchone()
    if not user or not check_password_hash(user['password_hash'], old_pw):
        conn.close()
        return jsonify({'error': 'Senha atual incorreta.'}), 403
    conn.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                 (generate_password_hash(new_pw), request.current_user['id']))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ── Analytics config ─────────────────────────────────────
@app.route('/api/analytics/config', methods=['GET'])
@auth_required
def get_analytics_config():
    env_saj = os.getenv('GA4_SERVICE_ACCOUNT_JSON', '').strip()
    env_pid = os.getenv('GA4_PROPERTY_ID', '').strip()
    env_mid = os.getenv('GA4_MEASUREMENT_ID', '').strip()
    conn = get_db()
    row = conn.execute('SELECT * FROM analytics_config ORDER BY id LIMIT 1').fetchone()
    conn.close()
    if not row:
        has_saj = bool(env_saj)
        pid = env_pid
        mid = env_mid
        return jsonify({
            'ga4_measurement_id': mid,
            'ga4_property_id': pid,
            'has_service_account': has_saj,
            'configured': bool(pid and has_saj),
        })
    data = dict(row)
    has_saj = bool(row['ga4_service_account_json']) or bool(env_saj)
    pid = data.get('ga4_property_id') or env_pid
    mid = data.get('ga4_measurement_id') or env_mid
    data['configured'] = bool(pid and has_saj)
    data['ga4_property_id'] = pid
    data['ga4_measurement_id'] = mid
    data.pop('ga4_service_account_json', None)  # never return the raw JSON to the frontend
    data['has_service_account'] = has_saj
    return jsonify(data)


@app.route('/api/analytics/config', methods=['PUT'])
@auth_required
def save_analytics_config():
    payload = request.get_json(force=True, silent=True) or {}
    mid = (payload.get('ga4_measurement_id') or '').strip()
    pid = (payload.get('ga4_property_id') or '').strip()
    saj = (payload.get('ga4_service_account_json') or '').strip()
    ts = now_iso()
    conn = get_db()
    existing = conn.execute('SELECT id, ga4_service_account_json FROM analytics_config ORDER BY id LIMIT 1').fetchone()
    if existing:
        # Keep old service account JSON if not provided
        final_saj = saj if saj else (existing['ga4_service_account_json'] or '')
        conn.execute(
            'UPDATE analytics_config SET ga4_measurement_id=?, ga4_property_id=?, ga4_service_account_json=?, updated_at=? WHERE id=?',
            (mid, pid, _encrypt(final_saj), ts, existing['id'])
        )
    else:
        conn.execute(
            'INSERT INTO analytics_config (ga4_measurement_id, ga4_property_id, ga4_service_account_json, updated_at) VALUES (?,?,?,?)',
            (mid, pid, _encrypt(saj), ts)
        )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ── GA4 Data API proxy ────────────────────────────────────
@app.route('/api/analytics/ga4', methods=['GET'])
@auth_required
def ga4_report():
    if not GA4_AVAILABLE:
        return jsonify({'error': 'google-analytics-data nao instalado no servidor.'}), 503

    # Prefer env vars, fall back to database
    property_id = os.getenv('GA4_PROPERTY_ID', '').strip()
    sa_json_str = os.getenv('GA4_SERVICE_ACCOUNT_JSON', '').strip()

    if not property_id or not sa_json_str:
        conn = get_db()
        row = conn.execute('SELECT * FROM analytics_config ORDER BY id LIMIT 1').fetchone()
        conn.close()
        if row:
            property_id = property_id or row['ga4_property_id']
            if not sa_json_str and row['ga4_service_account_json']:
                sa_json_str = _decrypt(row['ga4_service_account_json'])

    if not property_id or not sa_json_str:
        return jsonify({'configured': False}), 200

    try:
        sa_info = json.loads(sa_json_str)
        creds = ServiceAccountCredentials.from_service_account_info(
            sa_info,
            scopes=['https://www.googleapis.com/auth/analytics.readonly']
        )
        client = BetaAnalyticsDataClient(credentials=creds)

        # Overview: sessions, pageviews, users — last 30 days
        overview_req = RunReportRequest(
            property=f'properties/{property_id}',
            date_ranges=[DateRange(start_date='30daysAgo', end_date='today')],
            metrics=[
                Metric(name='sessions'),
                Metric(name='screenPageViews'),
                Metric(name='totalUsers'),
                Metric(name='averageSessionDuration'),
            ],
        )
        overview_resp = client.run_report(overview_req)
        row_ov = overview_resp.rows[0].metric_values if overview_resp.rows else []
        overview = {
            'sessions': int(row_ov[0].value) if row_ov else 0,
            'pageviews': int(row_ov[1].value) if row_ov else 0,
            'users': int(row_ov[2].value) if row_ov else 0,
            'avg_duration': round(float(row_ov[3].value)) if row_ov else 0,
        }

        # Daily pageviews — last 30 days
        daily_req = RunReportRequest(
            property=f'properties/{property_id}',
            date_ranges=[DateRange(start_date='30daysAgo', end_date='today')],
            dimensions=[Dimension(name='date')],
            metrics=[Metric(name='screenPageViews')],
            order_bys=[OrderBy(dimension=OrderBy.DimensionOrderBy(dimension_name='date'))],
        )
        daily_resp = client.run_report(daily_req)
        daily = [
            {'date': r.dimension_values[0].value, 'pageviews': int(r.metric_values[0].value)}
            for r in daily_resp.rows
        ]

        # Top 10 pages
        pages_req = RunReportRequest(
            property=f'properties/{property_id}',
            date_ranges=[DateRange(start_date='30daysAgo', end_date='today')],
            dimensions=[Dimension(name='pagePath')],
            metrics=[Metric(name='screenPageViews'), Metric(name='totalUsers')],
            order_bys=[OrderBy(metric=OrderBy.MetricOrderBy(metric_name='screenPageViews'), desc=True)],
            limit=10,
        )
        pages_resp = client.run_report(pages_req)
        top_pages = [
            {
                'path': r.dimension_values[0].value,
                'pageviews': int(r.metric_values[0].value),
                'users': int(r.metric_values[1].value),
            }
            for r in pages_resp.rows
        ]

        # Traffic sources
        sources_req = RunReportRequest(
            property=f'properties/{property_id}',
            date_ranges=[DateRange(start_date='30daysAgo', end_date='today')],
            dimensions=[Dimension(name='sessionDefaultChannelGroup')],
            metrics=[Metric(name='sessions')],
            order_bys=[OrderBy(metric=OrderBy.MetricOrderBy(metric_name='sessions'), desc=True)],
            limit=8,
        )
        sources_resp = client.run_report(sources_req)
        sources = [
            {'channel': r.dimension_values[0].value, 'sessions': int(r.metric_values[0].value)}
            for r in sources_resp.rows
        ]

        return jsonify({'configured': True, 'overview': overview, 'daily': daily, 'top_pages': top_pages, 'sources': sources})

    except Exception as exc:
        return jsonify({'error': str(exc), 'configured': True}), 500


# ── GitHub config ─────────────────────────────────────────
@app.route('/api/github/config', methods=['GET'])
@auth_required
def get_github_config():
    conn = get_db()
    row = conn.execute('SELECT * FROM github_config ORDER BY id LIMIT 1').fetchone()
    conn.close()
    if not row:
        return jsonify({'repo_owner': '', 'repo_name': '', 'branch': 'main', 'has_pat': False, 'configured': False})
    return jsonify({
        'repo_owner': row['repo_owner'],
        'repo_name': row['repo_name'],
        'branch': row['branch'],
        'has_pat': bool(row['github_pat']),
        'configured': bool(row['repo_owner'] and row['repo_name'] and row['github_pat']),
    })


@app.route('/api/github/config', methods=['PUT'])
@auth_required
def save_github_config():
    payload = request.get_json(force=True, silent=True) or {}
    owner = (payload.get('repo_owner') or '').strip()
    repo = (payload.get('repo_name') or '').strip()
    branch = (payload.get('branch') or 'main').strip()
    pat = (payload.get('github_pat') or '').strip()
    ts = now_iso()
    conn = get_db()
    existing = conn.execute('SELECT id, github_pat FROM github_config ORDER BY id LIMIT 1').fetchone()
    if existing:
        final_pat = pat if pat else (existing['github_pat'] or '')
        conn.execute(
            'UPDATE github_config SET repo_owner=?, repo_name=?, branch=?, github_pat=?, updated_at=? WHERE id=?',
            (owner, repo, branch, _encrypt(final_pat), ts, existing['id'])
        )
    else:
        conn.execute(
            'INSERT INTO github_config (repo_owner, repo_name, branch, github_pat, updated_at) VALUES (?,?,?,?,?)',
            (owner, repo, branch, _encrypt(pat), ts)
        )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ── GitHub API proxy ──────────────────────────────────────
def _safe_path(path_str):
    """Reject paths with traversal attempts or absolute paths.
    Allows an optional query string (after ?) with safe characters."""
    clean = (path_str or '').strip()
    if not clean or clean.startswith('/'):
        return None
    # Split path from optional query string
    parts = clean.split('?', 1)
    path_only = parts[0]
    if '..' in path_only:
        return None
    # Validate path portion: letters, digits, -, _, ., /
    if not re.match(r'^[\w\-./]+$', path_only):
        return None
    # Validate query string if present: allow word chars, =, &, %, +, -
    if len(parts) > 1 and not re.match(r'^[\w\-=&%.+]+$', parts[1]):
        return None
    return clean


@app.route('/api/github/proxy', methods=['POST'])
@auth_required
def github_proxy():
    payload = request.get_json(force=True, silent=True) or {}
    method = (payload.get('method') or 'GET').upper()
    path = _safe_path(payload.get('path') or '')
    body = payload.get('body')  # dict or None

    if method not in {'GET', 'PUT', 'DELETE'}:
        return jsonify({'error': 'Metodo nao permitido.'}), 405
    if not path:
        return jsonify({'error': 'Caminho invalido.'}), 400

    conn = get_db()
    row = conn.execute('SELECT * FROM github_config ORDER BY id LIMIT 1').fetchone()
    conn.close()
    if not row or not row['github_pat']:
        return jsonify({'error': 'GitHub nao configurado.'}), 503

    pat = _decrypt(row['github_pat'])
    owner = row['repo_owner']
    repo_name = row['repo_name']

    url = f'https://api.github.com/repos/{owner}/{repo_name}/{path}'
    headers = {
        'Authorization': f'token {pat}',
        'Accept': 'application/vnd.github+json',
        'X-GitHub-Api-Version': '2022-11-28',
        'User-Agent': 'iDialog-Admin/2.0',
    }

    try:
        req_body = json.dumps(body).encode('utf-8') if body else None
        if req_body:
            headers['Content-Type'] = 'application/json'

        gh_req = urllib.request.Request(url, data=req_body, headers=headers, method=method)
        with urllib.request.urlopen(gh_req, timeout=15) as resp:
            resp_body = resp.read()
            resp_data = json.loads(resp_body) if resp_body else {}
            return jsonify(resp_data), resp.status
    except urllib.error.HTTPError as exc:
        error_body = exc.read().decode('utf-8', errors='replace')
        try:
            error_data = json.loads(error_body)
        except Exception:
            error_data = {'message': error_body}
        return jsonify(error_data), exc.code
    except Exception as exc:
        return jsonify({'error': str(exc)}), 502


# ── Media library ─────────────────────────────────────────
@app.route('/api/media', methods=['GET'])
@auth_required
def list_media():
    files = []
    for fp in sorted(UPLOADS_DIR.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not fp.is_file():
            continue
        stat = fp.stat()
        ext = fp.suffix.lower().lstrip('.')
        mime_type = 'image' if ext in {'png', 'jpg', 'jpeg', 'webp', 'gif', 'svg'} else 'video' if ext in {'mp4', 'webm', 'mov'} else 'file'
        files.append({
            'filename': fp.name,
            'url': f"{request.host_url.rstrip('/')}/uploads/{fp.name}",
            'size': stat.st_size,
            'mime_type': mime_type,
            'modified': dt.datetime.fromtimestamp(stat.st_mtime).isoformat(),
        })
    return jsonify({'files': files})


@app.route('/api/media/<path:filename>', methods=['DELETE'])
@auth_required
def delete_media(filename):
    # Validate filename — no path traversal
    safe_name = secure_filename(filename)
    if not safe_name or safe_name != filename:
        return jsonify({'error': 'Nome de arquivo invalido.'}), 400
    fp = UPLOADS_DIR / safe_name
    if not fp.exists() or not fp.is_file():
        return jsonify({'error': 'Arquivo nao encontrado.'}), 404
    fp.unlink()
    return jsonify({'ok': True})


# ── Page blocks (public) ──────────────────────────────────
@app.route('/api/public/pages/<page_id>/blocks', methods=['GET'])
def public_page_blocks(page_id):
    clean_id = re.sub(r'[^\w\-]', '', page_id)
    conn = get_db()
    rows = conn.execute('SELECT block_key, content_html FROM page_blocks WHERE page_id = ?', (clean_id,)).fetchall()
    conn.close()
    return jsonify({'blocks': {row['block_key']: row['content_html'] for row in rows}})


# ── Page blocks (admin) ───────────────────────────────────
@app.route('/api/pages/<page_id>/blocks', methods=['GET'])
@auth_required
def list_page_blocks(page_id):
    clean_id = re.sub(r'[^\w\-]', '', page_id)
    conn = get_db()
    rows = conn.execute('SELECT block_key, content_html, updated_at FROM page_blocks WHERE page_id = ?', (clean_id,)).fetchall()
    conn.close()
    return jsonify({'blocks': [dict(r) for r in rows]})


@app.route('/api/pages/<page_id>/blocks/<block_key>', methods=['PUT'])
@auth_required
def save_page_block(page_id, block_key):
    clean_id = re.sub(r'[^\w\-]', '', page_id)
    clean_key = re.sub(r'[^\w\-]', '', block_key)
    payload = request.get_json(force=True, silent=True) or {}
    content_html = payload.get('content_html') or ''
    ts = now_iso()
    conn = get_db()
    conn.execute(
        'INSERT INTO page_blocks (page_id, block_key, content_html, updated_at) VALUES (?,?,?,?) '
        'ON CONFLICT(page_id, block_key) DO UPDATE SET content_html=excluded.content_html, updated_at=excluded.updated_at',
        (clean_id, clean_key, content_html, ts)
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


if __name__ == '__main__':
    port = int(os.getenv('PORT', '5001'))
    debug_mode = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug_mode)
